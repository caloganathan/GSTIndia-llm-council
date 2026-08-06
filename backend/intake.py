"""Notice ingestion.

A partner will not retype a notice into a form. Until the notice can be
uploaded, nothing else about this product matters.

The design principle here is that a GST notice is a highly structured
document, so most of the intake form can be filled WITHOUT a model at all:

    GSTIN            fixed 15-character format
    State            derived from the first two digits of the GSTIN
    Notice type      the form code appears literally in the document
    Reference number labelled and formatted
    Dates            a handful of Indian conventions
    Amounts          labelled with Rs./INR
    Section invoked  "section 73", "u/s 61"
    Tax period       "FY 2019-20", "2019-20"

Regex on a known format beats a model on every axis that matters here: it is
free, it is deterministic, it cannot hallucinate a GSTIN, and — the point that
decides it — nothing has to leave the machine.

A model is used for exactly the two fields regex cannot do: what the issues
are, and a summary of the facts. On the draft tier even that text is
anonymised first, so an uploaded notice is no less private than a typed one.

Scanned notices with no text layer are read by OCR where the optional engine
is installed (see `ocr.py`), and reported honestly where it is not. The
premise that a wrong OCR read is worse than an empty field still holds and is
enforced rather than avoided: OCR-derived text is marked as such, every field
read from it carries an `-ocr` source, and the review UI puts those fields in
the same must-confirm state as an unread amount.
"""

import io
import os
import re
from typing import Any, Dict, List, Optional, Tuple

from . import config, defects, notice_tables, ocr, sanitizer
from .openrouter import query_model

MAX_UPLOAD_BYTES = int(os.getenv("MAX_UPLOAD_BYTES", str(10 * 1024 * 1024)))
SUPPORTED = (".pdf", ".docx", ".txt")

# Text below this length almost certainly means a scanned page with no text
# layer rather than a genuinely short notice.
MIN_USEFUL_TEXT = 200

# How much of the notice reaches the model.
#
# This was 12,000, on the reasoning that the operative part of a notice is at
# the front. It is not. A real scrutiny attachment ran to 21,000 characters
# with its last three defects — including the only limb that went on to a show
# cause notice — beyond the cut, and they were never read at all. Defect
# segmentation now happens locally over the whole document, so the model sees
# the notice for context rather than for structure, but it still must see the
# operative part in full.
MAX_MODEL_CHARS = int(os.getenv("MAX_NOTICE_MODEL_CHARS", "48000"))


# ---------------------------------------------------------------------------
# GSTIN state codes — the first two digits identify the State, which drives
# the jurisdiction weighting. Deriving it costs nothing and is more reliable
# than asking anyone to pick from a dropdown of 36.
# ---------------------------------------------------------------------------

GSTIN_STATE_CODES = {
    "01": "Jammu and Kashmir", "02": "Himachal Pradesh", "03": "Punjab",
    "04": "Chandigarh", "05": "Uttarakhand", "06": "Haryana", "07": "Delhi",
    "08": "Rajasthan", "09": "Uttar Pradesh", "10": "Bihar", "11": "Sikkim",
    "12": "Arunachal Pradesh", "13": "Nagaland", "14": "Manipur",
    "15": "Mizoram", "16": "Tripura", "17": "Meghalaya", "18": "Assam",
    "19": "West Bengal", "20": "Jharkhand", "21": "Odisha",
    "22": "Chhattisgarh", "23": "Madhya Pradesh", "24": "Gujarat",
    "26": "Dadra and Nagar Haveli and Daman and Diu", "27": "Maharashtra",
    "29": "Karnataka", "30": "Goa", "31": "Lakshadweep", "32": "Kerala",
    "33": "Tamil Nadu", "34": "Puducherry",
    "35": "Andaman and Nicobar Islands", "36": "Telangana",
    "37": "Andhra Pradesh", "38": "Ladakh",
}

GSTIN_RE = re.compile(r"\b(\d{2})[A-Z]{5}\d{4}[A-Z][0-9A-Z]{3}\b")

# Portal-issued identifiers. These are the handles the officer files the matter
# under, and a reply that does not quote them can be rejected on its face.
#
# The shape is fixed: two letters, a run of digits, then a check character —
# ZD330226255583F for a scrutiny reference, AD330226081676X for an ARN. The
# first version of this matched on the WORD "notice" followed by capitals and
# duly extracted the reference "proposing" from a sentence, so the anchor is
# now the identifier's own format rather than the label beside it.
PORTAL_ID_RE = re.compile(r"\b([A-Z]{2}\d{10,14}[A-Z0-9])\b")

REFERENCE_RE = re.compile(
    r"(?:reference|ref|scrutiny\s+ref)\.?\s*(?:no|number)?\.?\s*[:\-]?\s*"
    r"([A-Z]{2}\d{10,14}[A-Z0-9])",
    re.IGNORECASE,
)

ARN_RE = re.compile(
    r"\bARN\s*[:\-]?\s*([A-Z]{2}\d{10,14}[A-Z0-9])", re.IGNORECASE
)

# Document Identification Number. Circular 128/2/2020-GST makes it mandatory,
# and its absence is a live ground of challenge — so it is captured whether or
# not the notice bothers to supply it.
DIN_RE = re.compile(r"\bDIN\s*[:\-]?\s*([A-Z0-9]{15,25})\b", re.IGNORECASE)

# 14.06.2026 / 14-06-2026 / 14/06/2026
NUMERIC_DATE_RE = re.compile(r"\b(\d{1,2})[./\-](\d{1,2})[./\-](\d{4})\b")
# 14 June 2026 / 14th June, 2026
TEXT_DATE_RE = re.compile(
    r"\b(\d{1,2})(?:st|nd|rd|th)?\s+"
    r"(January|February|March|April|May|June|July|August|September|October|"
    r"November|December)\,?\s+(\d{4})\b",
    re.IGNORECASE,
)
MONTHS = {m.lower(): i for i, m in enumerate(
    ["January", "February", "March", "April", "May", "June", "July",
     "August", "September", "October", "November", "December"], start=1)}

AMOUNT_RE = re.compile(
    r"(?:Rs\.?|INR|₹)\s*([\d,]+(?:\.\d{1,2})?)", re.IGNORECASE
)

FY_RE = re.compile(
    r"(?:F\.?Y\.?|financial\s+year|tax\s+period)\s*[:\-]?\s*(\d{4})\s*[-–/]\s*(\d{2,4})",
    re.IGNORECASE,
)
BARE_FY_RE = re.compile(r"\b(20\d{2})\s*[-–]\s*(\d{2})\b")

# Statutory references, WITH their sub-sections.
#
# The earlier version captured "16" from "section 16(2)(aa)" and then took only
# the first match in the whole notice. On a real scrutiny notice that reported
# the single provision "9" — read out of the phrase "section 9(5), if
# applicable" in boilerplate — for a notice issued under section 61. The
# sub-section is not a detail here: 16(2)(aa) and 16(4) are different disputes
# with different defences, and "16" names neither of them.
SECTION_RE = re.compile(
    r"(?:section|sec\.?|u/s)\s*(\d{1,3}[A-Z]?(?:\s*\(\s*[0-9a-zA-Z]{1,3}\s*\))*)",
    re.IGNORECASE,
)

RULE_RE = re.compile(
    r"\brules?\s*(\d{1,3}[A-Z]?(?:\s*\(\s*[0-9a-zA-Z]{1,3}\s*\))*)",
    re.IGNORECASE,
)

# The portal form states the provision in a labelled row. Where that row is
# present it is authoritative and beats anything scraped from the prose.
LABELLED_SECTION_RE = re.compile(
    r"Section\s+under\s+which\s+(?:the\s+)?notice\s+is\s+issued\s*[:\-]?\s*"
    r"(\d{1,3}[A-Z]?)",
    re.IGNORECASE,
)

LABELLED_DUE_DATE_RE = re.compile(
    r"Date\s+by\s+which\s+(?:the\s+)?reply\s+(?:has\s+to\s+be|is\s+to\s+be|to\s+be)"
    r"\s+submitted\s*[:\-]?\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
    re.IGNORECASE,
)

# A deadline the notice states in words, as the State letter formats do:
# "...explain the reasons for the above discrepancy on or before 19.01.2026".
#
# This is not the positional guess that was removed — that took the LAST date
# anywhere in the document and read an invoice date out of an annexure. Here
# the date is anchored to a direction to answer, within the preceding sentence.
# A date with no such direction beside it is still not a deadline.
ON_OR_BEFORE_RE = re.compile(
    r"\bon\s+or\s+before\s+(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})", re.IGNORECASE
)

REPLY_DIRECTION_RE = re.compile(
    r"\b(?:repl(?:y|ies)|response|respond|explain|explanation|objection|"
    r"submit|furnish|show\s+cause)\b",
    re.IGNORECASE,
)

DIRECTION_LOOKBACK = 200

# Most scrutiny notices print no deadline at all. They give a period running
# from a date the department does not know either: "within 30 days of receipt
# of the notice". Recording that period is the honest reading — the field is
# not missing, it is relative, and the difference matters to a reviewer
# deciding what still has to be found before the matter can be diarised.
# The parentheses are not decoration: departments print "within (30) days from
# the date of receipt of this notice" at least as often as they print it plain.
REPLY_WINDOW_RE = re.compile(
    r"within\s+\(?\s*(\d{1,3})\s*\)?\s+days\s+(?:of|from)\s+(?:the\s+)?"
    r"(?:date\s+of\s+)?(?:receipt|service|issue|issuance)",
    re.IGNORECASE,
)

LABELLED_NOTICE_DATE_RE = re.compile(
    r"Reference\s*No\.?\s*[:\-]?\s*[A-Z0-9]+\s*Date\s*[:\-]?\s*"
    r"(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
    re.IGNORECASE,
)

# The issuing officer, from the portal form's signature block.
OFFICER_NAME_RE = re.compile(
    r"Signature\s*\n?\s*Name\s*[:\-]\s*([A-Za-z][A-Za-z.\s]{2,40}?)\s*\n",
    re.IGNORECASE,
)
OFFICER_DESIGNATION_RE = re.compile(
    r"Designation\s*[:\-]\s*([A-Za-z()/,.\s]{4,60}?)\s*\n", re.IGNORECASE
)
# Jurisdiction runs onto a second line in the portal form ("RAM NAGAR ,
# Coimbatore-\nII , COIMBATORE , Tamil Nadu"), so a single-line match truncates
# the circle name mid-word.
JURISDICTION_RE = re.compile(
    r"Jurisdiction\s*[:\-]\s*([^\n]{3,90}(?:\n[^\n:]{1,60})?)", re.IGNORECASE
)
CIRCLE_RE = re.compile(
    r"Circle\s*[:\-]\s*([A-Za-z0-9 \-]{3,40})", re.IGNORECASE
)

# Entity name, identified by its statutory suffix.
#
# This is extracted locally for two reasons, and the second matters more: it
# fills the client name field, AND it gives the sanitiser something to scrub.
# Without it the draft tier strips the GSTIN and PAN but leaves the company
# name standing in the notice text — which is exactly the leak the test suite
# caught.
# The suffix group is case-insensitive but the name itself is not. Notices
# print the taxpayer in full capitals — "GRAM ENVOSOLUTION PRIVATE LIMITED" —
# and a wholly case-sensitive pattern silently found nothing on every one of
# them, leaving the sanitiser with no company name to scrub. Making the WHOLE
# pattern case-insensitive is not the fix either: it then matches ordinary
# prose such as "the limited relief sought".
#
# Separators inside a name are [ \t]+, never \s+. A notice prints the taxpayer
# and its address on consecutive lines:
#
#     M/s. MALABAR MARINE FOODS PRIVATE LIMITED
#     WILLINGDON ISLAND, KOCHI - 682003
#
# and \s+ crosses the newline, so the name came out as "MALABAR MARINE FOODS
# PRIVATE LIMITED\nWILLINGDON ISLAND". That reached the letterhead of the
# filing document, and it is what the sanitiser is given to scrub — so the
# scrub was also being aimed at the wrong string. An entity name never spans
# a line break on these forms.
ENTITY_RE = re.compile(
    r"\b((?:[A-Z][\w&.'\-]*[ \t]+){1,6}"
    r"(?i:PRIVATE[ \t]+LIMITED|PVT\.?[ \t]*LTD\.?|PUBLIC[ \t]+LIMITED|LIMITED"
    r"|LTD\.?|LLP|LIMITED[ \t]+LIABILITY[ \t]+PARTNERSHIP|&[ \t]*CO\.?"
    r"|AND[ \t]+COMPANY|&[ \t]*SONS|ENTERPRISES|INDUSTRIES|ASSOCIATES"
    r"|TRADERS|AGENCIES|ELECTRICALS|HOTELS))\b"
)

# "Tvl." (Tamil Nadu), "M/s." and "Messrs." introduce the taxpayer by name and
# are the most reliable signal of all where present. The prefix may sit at the
# end of its own line, so one newline is allowed between prefix and name — but
# none inside the name itself.
PREFIXED_ENTITY_RE = re.compile(
    r"(?:M/s\.?|Tvl\.?|Messrs\.?)[ \t]*\n?[ \t]*"
    r"([A-Z][\w&.'\-]*(?:[ \t]+[A-Z][\w&.'\-]*){0,7})"
)

DUE_DATE_HINT_RE = re.compile(
    r"(?:reply|response|submit|furnish|show cause)[^.]{0,120}?"
    r"(?:on or before|within|by)\s+([^.\n]{0,60})",
    re.IGNORECASE,
)

# The taxpayer's own name, from the labelled block rather than from anywhere
# in the document.
#
# THIS IS THE FIELD THE ANNEXURE STEALS. A notice carries a supplier-wise ITC
# annexure listing dozens of counterparties, every one of them a company with
# a statutory suffix, and a document-wide search for that suffix returns
# whichever supplier happens to have the longest name. On a real corpus that
# put "THE NEW INDIA ASSURANCE CO LIMITED" — row 42 of the annexure, with its
# own GSTIN — on the letterhead of a reply filed for a packaging firm whose
# name was printed, correctly and plainly, on page 1. Eight of thirteen
# notices came back with the wrong entity, and none of them was flagged,
# because a match had been found.
#
# The name is therefore taken from where the form states it, in this order:
# the labelled row, then the taxpayer block, then a prefix near the taxpayer's
# own GSTIN. The document-wide search survives only as the last resort, for
# text that carries no GSTIN at all.

# "Legal Name: Tvl. F Care Plus Llp" — the letter format states it inline.
LABELLED_ENTITY_RE = re.compile(
    r"(?:Legal|Trade)\s*Name\s*[:\-]\s*(?:M/s\.?|Tvl\.?|Messrs\.?)?[ \t]*"
    r"([A-Za-z][\w&.'\-]*(?:[ \t]+[A-Za-z][\w&.'\-]*){0,7})"
)

# The portal attachment prints a taxpayer block instead: three labels and
# three values, with the GSTIN last. Its furniture is what identifies the
# block as the taxpayer's rather than a supplier row in an annexure.
TAXPAYER_BLOCK_RE = re.compile(
    r"Details\s+of\s+the\s+Tax\s?payer|Reg\s+Status|Trade\s+Name",
    re.IGNORECASE,
)

# How far from the taxpayer's GSTIN the name is still the taxpayer's name.
ENTITY_WINDOW = 600

# Lines inside the taxpayer block that are labels, column rulers or status
# values rather than the name.
_NOT_A_NAME = {
    "gstin", "gstn", "pan", "name", "trade name", "legal name", "reg status",
    "status", "active", "inactive", "financial year", "zone", "circle",
    "details of the tax payer", "details of the taxpayer", "office details",
    "designation of the proper officer",
}


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def extract_text(filename: str, content: bytes) -> Tuple[str, List[str]]:
    """
    Pull text out of an uploaded notice. Entirely local — nothing is sent
    anywhere at this stage.

    Returns (text, warnings). Warnings are shown to the user rather than
    swallowed: an empty field they can see is safe, one they cannot is not.

    Thin wrapper over `extract_document`, kept because most callers want only
    the text and the warnings.
    """
    document = extract_document(filename, content)
    return document["text"], document["warnings"]


def extract_document(filename: str, content: bytes) -> Dict[str, Any]:
    """
    Read an uploaded notice and report HOW it was read.

    The provenance is not bookkeeping. Text lifted from a PDF's text layer is
    the document; text recovered by OCR from an image of the document is a
    machine's reading of it, and every figure in it is a proposal rather than
    a fact. Callers propagate that distinction all the way to the review UI,
    so nothing that came out of a scanner is ever presented as confirmed.

    Returns {text, warnings, source, ocr}, where `source` is one of
    "text-layer", "ocr" or "none".
    """
    warnings: List[str] = []
    name = (filename or "").lower()

    if len(content) > MAX_UPLOAD_BYTES:
        raise ValueError(
            f"File is larger than {MAX_UPLOAD_BYTES // (1024 * 1024)} MB."
        )
    if not any(name.endswith(ext) for ext in SUPPORTED):
        raise ValueError(
            f"Unsupported file type. Upload one of: {', '.join(SUPPORTED)}."
        )

    text = ""
    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    raise ValueError(
                        "This PDF is password protected. Remove the password "
                        "and upload again."
                    )
            pages = []
            for page in reader.pages:
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    warnings.append("One page could not be read and was skipped.")
            text = "\n".join(pages)
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Could not read the PDF: {e}")

    elif name.endswith(".docx"):
        try:
            from docx import Document
            document = Document(io.BytesIO(content))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            text = "\n".join(parts)
        except Exception as e:
            raise ValueError(f"Could not read the document: {e}")

    else:
        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

    text = _tidy(text)
    source = "text-layer" if text else "none"
    ocr_result = None

    # A PDF that yields almost nothing is a scan. Try to read the image before
    # giving up on it — but only for PDFs, and only when the text layer has
    # genuinely failed. A notice that already has text is never re-read by
    # OCR: the text layer is the document, and OCR is a reading of a picture
    # of it.
    if len(text) < MIN_USEFUL_TEXT and name.endswith(".pdf"):
        engine_ready, reason = ocr.available()
        if engine_ready:
            try:
                ocr_result = ocr.ocr_pdf(content)
            except Exception as e:
                warnings.append(
                    f"This notice appears to be scanned and OCR failed on it "
                    f"({e}). Paste the text of the notice manually."
                )
            else:
                recovered = _tidy(ocr_result.get("text", ""))
                if len(recovered) > len(text):
                    text = recovered
                    source = "ocr"
                warnings.extend(ocr.describe_quality(ocr_result))
        else:
            warnings.append(
                "This notice appears to be a scan with no text layer. "
                f"{reason} Until then, paste the text of the notice manually."
            )

    if len(text) < MIN_USEFUL_TEXT and source != "ocr":
        warnings.append(
            "Very little text could be read. This is usually a scanned notice "
            "with no text layer — paste the text of the notice manually, or "
            "upload a digitally generated copy."
        )

    return {
        "text": text,
        "warnings": warnings,
        "source": source,
        "ocr": ocr_result,
    }


def _tidy(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text or "")
    return re.sub(r"\n{3,}", "\n\n", text).strip()


# ---------------------------------------------------------------------------
# Local field extraction — no model, nothing leaves the machine
# ---------------------------------------------------------------------------


def _iso_date(day: int, month: int, year: int) -> Optional[str]:
    if not (1 <= month <= 12 and 1 <= day <= 31 and 1900 < year < 2200):
        return None
    return f"{year:04d}-{month:02d}-{day:02d}"


def find_dates(text: str) -> List[str]:
    """All dates in the notice, in document order, as ISO strings."""
    found: List[Tuple[int, str]] = []
    for match in NUMERIC_DATE_RE.finditer(text):
        day, month, year = (int(g) for g in match.groups())
        iso = _iso_date(day, month, year)
        if iso:
            found.append((match.start(), iso))
    for match in TEXT_DATE_RE.finditer(text):
        day, month_name, year = match.groups()
        iso = _iso_date(int(day), MONTHS[month_name.lower()], int(year))
        if iso:
            found.append((match.start(), iso))
    found.sort()
    seen, ordered = set(), []
    for _, iso in found:
        if iso not in seen:
            seen.add(iso)
            ordered.append(iso)
    return ordered


def find_notice_type(text: str, pack, with_offset: bool = False):
    """
    Identify the form code. Longest codes are matched first so that DRC-01A is
    never mistaken for DRC-01.
    """
    codes = sorted(
        (c for c in pack.NOTICE_TYPES if c != "OTHER"),
        key=len, reverse=True,
    )
    upper = text.upper()
    for code in codes:
        # Tolerate ASMT-10, ASMT 10, ASMT10 and "GST ASMT - 10", which is how
        # the portal actually prints the form header.
        pattern = code.replace("-", r"[\s\-]*")
        match = re.search(rf"\b{pattern}\b", upper)
        if match:
            return (code, match.start()) if with_offset else code
    return (None, None) if with_offset else None


def find_sections(text: str) -> List[str]:
    """
    Every statutory section in the text, sub-sections intact, in document order.

    Deduplicated but never truncated to one: a notice engages several
    provisions and the reply has to answer each of them.
    """
    return _collect(SECTION_RE, text)


def find_rules(text: str) -> List[str]:
    return _collect(RULE_RE, text)


def _collect(pattern: re.Pattern, text: str) -> List[str]:
    seen, ordered = set(), []
    for match in pattern.finditer(text or ""):
        # "16 ( 2 ) ( aa )" and "16(2)(aa)" are the same provision.
        reference = re.sub(r"\s+", "", match.group(1))
        if reference and reference not in seen:
            seen.add(reference)
            ordered.append(reference)
    return ordered


def _looks_like_a_name(line: str) -> bool:
    """Whether a line inside the taxpayer block is the name rather than a label."""
    cleaned = line.strip().strip(".,;")
    if not (3 <= len(cleaned) <= 80):
        return False
    if ":" in cleaned or cleaned[0].isdigit():
        return False
    if not re.search(r"[A-Za-z]{2}", cleaned):
        return False
    return re.sub(r"\s+", " ", cleaned).lower() not in _NOT_A_NAME


def _name_from_taxpayer_block(text: str, gstin: re.Match):
    """
    The name printed directly above the taxpayer's own GSTIN.

    The block prints Trade Name, Legal Name and GSTIN as three consecutive
    values, so the line above the GSTIN is the legal name. Two conditions keep
    this off the supplier rows of an annexure, where a name also sits beside a
    GSTIN: the GSTIN must stand alone on its line, which a tabulated annexure
    row never does, and the block's own furniture must be nearby.
    """
    line_start = text.rfind("\n", 0, gstin.start()) + 1
    line_end = text.find("\n", gstin.end())
    line_end = len(text) if line_end == -1 else line_end
    if text[line_start:line_end].strip() != gstin.group(0):
        return None, None

    around = text[max(0, gstin.start() - ENTITY_WINDOW):
                  gstin.end() + ENTITY_WINDOW]
    if not TAXPAYER_BLOCK_RE.search(around):
        return None, None

    offset = line_start
    for raw in reversed(text[:line_start].splitlines()[-3:]):
        offset -= len(raw) + 1
        if not raw.strip():
            continue
        if _looks_like_a_name(raw):
            return raw.strip().strip(".,;"), max(offset, 0)
        break
    return None, None


def _best_in(matches, floor: int = 0):
    """The longest capture in a list of matches, with its offset."""
    candidates = [(m.group(1).strip(" .,"), m.start()) for m in matches
                  if len(m.group(1).strip(" .,")) >= floor]
    if not candidates:
        return None, None
    return max(candidates, key=lambda pair: len(pair[0]))


def find_entity_name(text: str, with_offset: bool = False):
    """
    The taxpayer's name, read from where the form states it.

    In descending order of authority: the labelled row, the taxpayer block
    above the GSTIN, a "Tvl."/"M/s." prefix near the taxpayer's own GSTIN, and
    a statutory suffix near it. Only when the text carries no GSTIN at all does
    the search widen to the whole document — see LABELLED_ENTITY_RE above for
    what a document-wide search does to a notice with a supplier annexure.
    """
    text = text or ""

    labelled = LABELLED_ENTITY_RE.search(text)
    if labelled:
        name = labelled.group(1).strip(" .,")
        return (name, labelled.start(1)) if with_offset else name

    gstin = GSTIN_RE.search(text)
    if gstin:
        name, at = _name_from_taxpayer_block(text, gstin)
        if name:
            return (name, at) if with_offset else name

        start = max(0, gstin.start() - ENTITY_WINDOW)
        window = text[start:gstin.end() + ENTITY_WINDOW]
        for pattern in (PREFIXED_ENTITY_RE, ENTITY_RE):
            name, at = _best_in(pattern.finditer(window))
            if name:
                return (name, at + start) if with_offset else name

    for pattern in (PREFIXED_ENTITY_RE, ENTITY_RE):
        name, at = _best_in(pattern.finditer(text))
        if name:
            return (name, at) if with_offset else name

    return (None, None) if with_offset else None


def find_amounts(text: str) -> List[float]:
    amounts = []
    for match in AMOUNT_RE.finditer(text):
        try:
            amounts.append(float(match.group(1).replace(",", "")))
        except ValueError:
            continue
    return amounts


def find_tax_period(text: str, with_offset: bool = False):
    match = FY_RE.search(text)
    if not match:
        match = BARE_FY_RE.search(text)
    if not match:
        return (None, None) if with_offset else None
    start, end = match.group(1), match.group(2)
    if len(end) == 4:
        end = end[2:]
    period = f"FY {start}-{end}"
    return (period, match.start()) if with_offset else period


def _directed_deadline(text: str) -> Optional[re.Match]:
    """
    An "on or before <date>" that the notice attaches to a direction to answer.

    The direction must appear in the run-up to the date, so a date standing on
    its own — an invoice date in an annexure, a rate notification — is not
    promoted to a deadline merely by being formatted like one.
    """
    for match in ON_OR_BEFORE_RE.finditer(text or ""):
        lead = text[max(0, match.start() - DIRECTION_LOOKBACK):match.start()]
        if REPLY_DIRECTION_RE.search(lead):
            return match
    return None


def _iso_from_numeric(raw: str) -> Optional[str]:
    match = NUMERIC_DATE_RE.search(raw or "")
    if not match:
        return None
    day, month, year = (int(g) for g in match.groups())
    return _iso_date(day, month, year)


def extract_fields_local(text: str, pack) -> Dict[str, Any]:
    """
    Everything derivable without a model.

    Each field carries where it came from, so the UI can show the user what was
    read off the notice versus what a model proposed.

    Dates and the provision invoked are taken from the portal form's LABELLED
    rows wherever those rows exist, and only fall back to position heuristics
    when they do not. The heuristics that used to run unconditionally — first
    date in the document is the notice date, last date is the deadline — read a
    rate-notification date from 2022 as the date of a 2026 notice and an
    invoice date from an annexure as the reply deadline. Position is a poor
    proxy for meaning in a document that carries dozens of dates.
    """
    fields: Dict[str, Any] = {}
    sources: Dict[str, str] = {}
    snippets: Dict[str, Dict[str, Any]] = {}

    def put(key: str, value: Any, source: str, at: Any = None):
        if value not in (None, "", []):
            fields[key] = value
            sources[key] = source
            offset = _offset_of(at)
            if offset is not None:
                snippets[key] = snippet_at(text, offset)

    gstin_match = GSTIN_RE.search(text)
    if gstin_match:
        put("gstin", gstin_match.group(0), "notice", gstin_match)
        put("state", GSTIN_STATE_CODES.get(gstin_match.group(1)), "gstin",
            gstin_match)

    name, name_at = find_entity_name(text, with_offset=True)
    put("client_name", name, "notice", name_at)
    notice_type, type_at = find_notice_type(text, pack, with_offset=True)
    put("notice_type", notice_type, "notice", type_at)
    period, period_at = find_tax_period(text, with_offset=True)
    put("tax_period", period, "notice", period_at)

    # --- Identifiers ------------------------------------------------------
    reference = REFERENCE_RE.search(text)
    if reference:
        put("notice_reference", reference.group(1).strip(), "notice", reference)
    else:
        loose = PORTAL_ID_RE.search(text)
        if loose:
            put("notice_reference", loose.group(1), "notice", loose)

    arn = ARN_RE.search(text)
    if arn:
        put("notice_arn", arn.group(1), "notice", arn)

    din = DIN_RE.search(text)
    if din:
        put("din", din.group(1), "notice", din)

    # --- Provision invoked ------------------------------------------------
    labelled_section = LABELLED_SECTION_RE.search(text)
    sections = find_sections(text)
    if labelled_section:
        put("section_invoked", labelled_section.group(1), "notice-labelled",
            labelled_section)
    elif sections:
        put("section_invoked", sections[0], "notice", SECTION_RE.search(text))
    put("sections_cited", sections, "notice")
    put("rules_cited", find_rules(text), "notice")

    # --- Dates ------------------------------------------------------------
    labelled_notice_date = LABELLED_NOTICE_DATE_RE.search(text)
    if labelled_notice_date:
        put("notice_date", _iso_from_numeric(labelled_notice_date.group(1)),
            "notice-labelled", labelled_notice_date)

    labelled_due = LABELLED_DUE_DATE_RE.search(text)
    if labelled_due:
        put("due_date", _iso_from_numeric(labelled_due.group(1)),
            "notice-labelled", labelled_due)
    else:
        directed = _directed_deadline(text)
        if directed:
            put("due_date", _iso_from_numeric(directed.group(1)),
                "notice-directed", directed)

    window = REPLY_WINDOW_RE.search(text)
    if window and not fields.get("due_date"):
        put("reply_window_days", int(window.group(1)), "notice", window)

    if not fields.get("notice_date"):
        dates = find_dates(text)
        if dates:
            put("notice_date", dates[0], "notice-inferred",
                _first_date_offset(text, dates[0]))

    # --- Issuing officer --------------------------------------------------
    officer_name = OFFICER_NAME_RE.search(text)
    designation = OFFICER_DESIGNATION_RE.search(text)
    parts = [
        officer_name.group(1).strip() if officer_name else "",
        designation.group(1).strip() if designation else "",
    ]
    put("issuing_officer", ", ".join(p for p in parts if p), "notice",
        officer_name or designation)

    jurisdiction = JURISDICTION_RE.search(text)
    if jurisdiction:
        office = re.sub(r"\s*\n\s*", "", jurisdiction.group(1))
        office = re.sub(r"\s*,\s*", ", ", office).strip(" ,")
        put("jurisdiction_office", office, "notice", jurisdiction)
    else:
        circle = CIRCLE_RE.search(text)
        if circle:
            put("jurisdiction_office", circle.group(1).strip(" ,"), "notice",
                circle)

    return {"fields": fields, "sources": sources, "snippets": snippets}


# ---------------------------------------------------------------------------
# Provenance snippets — the reviewer's shortcut back to the notice
# ---------------------------------------------------------------------------
# Checking extraction is the slowest step in using this product, and it is the
# step that decides whether a firm trusts it. Reading a field off a form and
# then hunting through twenty pages of PDF for the sentence it came from is
# what makes that step slow. So every locally extracted field carries the text
# around it: the reviewer confirms or corrects in place, without opening the
# notice at all.
#
# Only local extraction produces snippets. A model-proposed field has no
# offset in the document because it was not read off a position — it was
# summarised from the whole — and inventing one would be a lie about
# provenance in a product whose whole argument is that it does not do that.

SNIPPET_BEFORE = 90
SNIPPET_AFTER = 110


def snippet_at(text: str, offset: int) -> Dict[str, Any]:
    """The text around an offset, snapped to word boundaries."""
    if offset is None or offset < 0 or not text:
        return {}
    start = max(0, offset - SNIPPET_BEFORE)
    end = min(len(text), offset + SNIPPET_AFTER)

    # Snap outward to whitespace so the excerpt does not begin or end
    # mid-word, which reads as corruption rather than as an extract.
    if start > 0:
        space = text.find(" ", start, offset)
        if space != -1:
            start = space + 1
    if end < len(text):
        space = text.rfind(" ", offset, end)
        if space != -1:
            end = space

    excerpt = re.sub(r"\s+", " ", text[start:end]).strip()
    return {
        "text": excerpt,
        "offset": offset,
        "truncated_start": start > 0,
        "truncated_end": end < len(text),
    }


def _offset_of(at: Any) -> Optional[int]:
    if at is None:
        return None
    if isinstance(at, int):
        return at
    start = getattr(at, "start", None)
    return start() if callable(start) else None


def _first_date_offset(text: str, iso: str) -> Optional[int]:
    """Where the date that produced this ISO string actually sits."""
    for match in NUMERIC_DATE_RE.finditer(text):
        day, month, year = (int(g) for g in match.groups())
        if _iso_date(day, month, year) == iso:
            return match.start()
    for match in TEXT_DATE_RE.finditer(text):
        day, month_name, year = match.groups()
        if _iso_date(int(day), MONTHS[month_name.lower()], int(year)) == iso:
            return match.start()
    return None


def extract_defects(text: str, pack) -> List[Dict[str, Any]]:
    """
    Decompose a notice into the limbs it will actually be decided in.

    Each defect gets its own head-wise amount read straight off the department's
    own annexure, its own sections, and the evidence list for its type. Where a
    figure cannot be read with confidence the defect is flagged rather than
    filled — a reviewer who can see an empty amount will supply it; one who
    cannot see a wrong amount will file it.
    """
    catalogue = getattr(pack, "DEFECT_TYPES", [])
    found = defects.segment(text, catalogue)
    if not found:
        return []

    head_order = notice_tables.detect_head_order(text)

    for defect in found:
        body = defect.get("notice_extract", "")
        row = notice_tables.read_defect_amount(
            body, head_order, defect.pop("preamble", ""),
        )
        if row:
            defect["amount_by_head"] = defects.normalise_heads(row["amounts"])
            defect["amount_basis"] = row.get("basis")
            defect["amount_label"] = row.get("label")
        else:
            defect["amount_unread"] = True

        # Sections stated inside the limb beat the catalogue's defaults, which
        # are only a starting point.
        cited = find_sections(body)
        if cited:
            defect["sections"] = _merge_preserving_order(defect["sections"], cited)
        cited_rules = find_rules(body)
        if cited_rules:
            defect["rules"] = _merge_preserving_order(defect["rules"], cited_rules)

    return found


def _merge_preserving_order(primary: List[str], extra: List[str]) -> List[str]:
    merged = list(primary)
    for item in extra:
        if item not in merged:
            merged.append(item)
    return merged


# ---------------------------------------------------------------------------
# Model-assisted extraction — only the two fields regex cannot do
# ---------------------------------------------------------------------------


def _build_reading_prompt(text: str, pack, notice_type: Optional[str]) -> str:
    notice = pack.NOTICE_TYPES.get(notice_type) if notice_type else None
    context = (
        f"The document appears to be a {notice.code} — {notice.name}, issued "
        f"under {notice.statute}.\n\n" if notice else ""
    )
    return f"""\
Read the {pack.SHORT_NAME} notice below and report only what it says. You are
transcribing, not advising: do not analyse the merits, do not suggest a reply,
and do not add anything the notice does not contain.

{context}NOTICE:
\"\"\"
{text[:MAX_MODEL_CHARS]}
\"\"\"

Return a SINGLE JSON object and nothing else:

{{
  "issues": "The discrepancies or allegations the notice raises, one per line, numbered. Use the notice's own framing and its own figures. If it annexes a table of differences, describe what the table alleges rather than reproducing it.",
  "facts": "Two to four sentences on what the notice states about the taxpayer and the period: nature of business if mentioned, what the officer says was found, and any prior correspondence referred to. Facts only, from the notice.",
  "officer": "Designation and office of the issuing officer, if stated, else empty",
  "confidence": "high" | "medium" | "low"
}}

Set confidence to "low" if the text is garbled, truncated, or does not read
like a tax notice. An honest "low" is useful; a confident misreading is not."""


async def extract_fields_assisted(
    text: str,
    pack,
    tier: Dict[str, Any],
    notice_type: Optional[str] = None,
) -> Tuple[Dict[str, Any], Optional[Dict[str, Any]], List[str]]:
    """
    Read the issues and facts out of the notice.

    On the anonymising tier the notice text is scrubbed before it is sent, so
    an uploaded notice is no less private than a typed one.
    """
    warnings: List[str] = []
    if not text.strip():
        return {}, None, ["No text to read."]

    outgoing = text
    # The replacement map is kept so the model's answer can be restored before
    # it reaches the user: a discarded map meant the extracted issues carried
    # "[GSTIN-1]" and "the Taxpayer" into the matter and then into both
    # exported documents. The model still never sees the real values.
    replacements: Dict[str, str] = {}
    if tier.get("anonymise"):
        # The entity name must be found before scrubbing, or the sanitiser has
        # nothing to match on and the company name goes out in clear.
        outgoing = sanitizer.scrub_text(text, replacements,
                                        client_name=find_entity_name(text))
        leaks = sanitizer.audit_leaks(outgoing[:MAX_MODEL_CHARS])
        if leaks:
            return {}, None, [
                "Identifiers could not be fully removed from this notice, so "
                f"it was not sent for reading ({', '.join(leaks)}). Enter the "
                "issues and facts manually, or use the Pro tier."
            ]

    model = tier.get("grounding") or tier.get("verifier") or ""
    if not model:
        return {}, None, ["No model configured for reading notices."]

    result = await query_model(
        model,
        [{"role": "user", "content": _build_reading_prompt(outgoing, pack, notice_type)}],
        effort=config.role_effort("briefing"),
        max_tokens=config.role_max_tokens("briefing"),
        zdr=config.ENFORCE_ZDR and not tier.get("anonymise"),
    )

    if not result.get("ok"):
        return {}, None, [
            f"The notice could not be read automatically ({result.get('error')}). "
            "Enter the issues and facts manually."
        ]

    parsed = _parse_json(result.get("content", ""))
    if parsed is None:
        return {}, result.get("usage"), [
            "The notice was read but the result could not be interpreted. "
            "Enter the issues and facts manually."
        ]

    if parsed.get("confidence") == "low":
        warnings.append(
            "The notice was difficult to read — check the issues and facts "
            "carefully before running the panel."
        )

    fields = {}
    for key in ("issues", "facts"):
        value = (parsed.get(key) or "").strip()
        if value:
            fields[key] = value
    if parsed.get("officer"):
        fields["issuing_officer"] = parsed["officer"].strip()

    if replacements:
        fields = sanitizer.restore_structure(fields, replacements)

    return fields, result.get("usage"), warnings


def _parse_json(text: str) -> Optional[Dict[str, Any]]:
    import json
    for candidate in _json_candidates(text):
        try:
            parsed = json.loads(candidate)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            continue
    return None


def _json_candidates(text: str):
    if not text:
        return
    yield text
    fence = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if fence:
        yield fence.group(1)
    start, end = text.find("{"), text.rfind("}")
    if start != -1 and end > start:
        yield text[start:end + 1]


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------


async def read_notice(
    filename: str,
    content: bytes,
    pack,
    tier: Dict[str, Any],
    use_model: bool = True,
) -> Dict[str, Any]:
    """Single-document convenience wrapper over `read_notice_set`."""
    return await read_notice_set([(filename, content)], pack, tier, use_model)


async def read_notice_set(
    documents: List[Tuple[str, bytes]],
    pack,
    tier: Dict[str, Any],
    use_model: bool = True,
) -> Dict[str, Any]:
    """
    Turn one or more uploaded documents into a proposed matter.

    A scrutiny notice does not arrive as one file. The portal issues a one-page
    form carrying the reference number, the provision, the reply date and the
    officer's name, and attaches a separate document — often twenty pages —
    carrying the actual defects and their annexures. Neither is sufficient
    alone, and asking a partner to upload them one at a time and reconcile the
    result by hand defeats the point of the upload.

    Documents are read in order and the FIRST value found for a field wins, so
    a portal form uploaded alongside its attachment supplies the identifiers
    while the attachment supplies the defects.

    Nothing here is authoritative. Everything is a proposal the user reviews
    before the panel runs — a wrong tax period silently accepted is worse than
    an empty one.
    """
    warnings: List[str] = []
    fields: Dict[str, Any] = {}
    sources: Dict[str, str] = {}
    snippets: Dict[str, Dict[str, Any]] = {}
    read_documents: List[Dict[str, Any]] = []
    combined: List[str] = []
    found_defects: List[Dict[str, Any]] = []
    usage = None
    any_scanned = False

    for filename, content in documents:
        document = extract_document(filename, content)
        text = document["text"]
        warnings.extend(f"{filename}: {w}" for w in document["warnings"])
        combined.append(text)
        scanned = document["source"] == "ocr"
        if scanned:
            any_scanned = True

        local = extract_fields_local(text, pack)
        for key, value in local["fields"].items():
            if key not in fields:
                fields[key] = value
                source = local["sources"].get(key, "notice")
                # An OCR-read field is a machine's reading of a picture of the
                # notice, and the review UI treats it accordingly. The suffix
                # is what carries that through — losing it here would make a
                # scanned figure indistinguishable from a printed one.
                sources[key] = f"{source}-ocr" if scanned else source
                snippet = local.get("snippets", {}).get(key)
                if snippet:
                    snippets[key] = {**snippet, "filename": filename,
                                     "scanned": scanned}

        # The document with the most defects is the attachment; a portal form
        # or a covering letter contributes none and must not displace it.
        document_defects = extract_defects(text, pack)
        if scanned:
            for defect in document_defects:
                defect["from_scan"] = True
        if len(document_defects) > len(found_defects):
            found_defects = document_defects

        read_documents.append({
            "filename": filename,
            "text_length": len(text),
            "defects_found": len(document_defects),
            "source": document["source"],
            "ocr": document["ocr"],
        })

    text = "\n\n".join(t for t in combined if t)

    if use_model and len(text) >= MIN_USEFUL_TEXT:
        assisted, usage, model_warnings = await extract_fields_assisted(
            text, pack, tier, fields.get("notice_type")
        )
        for key, value in assisted.items():
            fields[key] = value
            sources[key] = "read"
        warnings.extend(model_warnings)

    if found_defects:
        fields["defects"] = found_defects
        sources["defects"] = "notice-ocr" if any_scanned else "notice"
        summary = defects.triage(found_defects)
        unread = [d["heading"] for d in found_defects if d.get("amount_unread")]

        # With every limb read, the sum of the limbs is the better figure — it
        # is built from the department's own head-wise annexure. With a limb
        # UNREAD the sum understates, so a total the notice printed for itself
        # is preferred over our incomplete arithmetic. Overwriting it with the
        # short sum stated a definite, wrong, taxpayer-favourable number and
        # showed it to the reviewer as read "from defects".
        if unread and fields.get("amount_disputed"):
            fields["amount_disputed_from_limbs"] = summary["total_amount"]
            sources["amount_disputed_from_limbs"] = "defects-incomplete"
        else:
            fields["amount_disputed"] = summary["total_amount"]
            sources["amount_disputed"] = (
                "defects-incomplete" if unread else "defects")
        fields["amount_incomplete"] = bool(unread)

        if unread:
            warnings.append(
                "The amount could not be read for "
                f"{len(unread)} of {len(found_defects)} defects "
                f"({'; '.join(unread[:4])}"
                f"{'…' if len(unread) > 4 else ''}). Enter these from the "
                "notice annexure before running the panel."
            )

        scanned_with_figures = [d for d in found_defects
                                if d.get("from_scan") and not d.get("amount_unread")]
        if scanned_with_figures:
            warnings.append(
                f"{len(scanned_with_figures)} defect amount(s) were read by OCR "
                "from a scanned image. Check each against the notice annexure — "
                "these are the figures the reply will quote back to the officer."
            )
    else:
        warnings.append(
            "No defect headings were found, so this notice could not be broken "
            "into limbs. Add them by hand — a reply that answers a multi-limb "
            "notice as one issue concedes ground it need not concede."
        )

    missing = [f for f in ("notice_type", "state", "tax_period", "facts")
               if not fields.get(f)]
    if missing:
        warnings.append(
            "Could not determine: " + ", ".join(missing) +
            ". Complete these before running the panel."
        )

    return {
        "fields": fields,
        "sources": sources,
        # Where each locally extracted field came from, so the reviewer can
        # confirm it without opening the PDF.
        "snippets": snippets,
        "warnings": warnings,
        "documents": read_documents,
        "text_length": len(text),
        "scanned": any_scanned,
        "usage": usage,
        # The extracted text is returned so the user can see what was read and
        # copy from it. The uploaded files themselves are never persisted.
        "text": text[:MAX_MODEL_CHARS],
    }
