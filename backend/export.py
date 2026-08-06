"""Reply pack export — two documents, never one.

WHY TWO
-------
The first version of this module produced a single file containing both the
draft submission and the firm's internal analysis. That file was read by a
practising partner, who said it could not be sent to a client, let alone filed.
He was right, and the reason was worse than presentation.

Section 5 of that document was headed "Points for Reviewer Attention" and
contained, among other things, "worst realistic monetary exposure … approx.
Rs. 2.02 lakh" and "this is a real exposure and must be quantified, not assumed
away". Section 7 recorded that the firm's own confidence was "defensible, not
strong, because … the team has not yet verified line-by-line." Every word of
that is proper working-paper content. All of it sat in the same file as the
text intended for the proper officer.

So there are now two builders and no option to merge them:

    build_filing_reply()   goes to the department, over the CLIENT's
                           letterhead and its authorised signatory. Contains
                           submissions, figures, authorities that VERIFIED,
                           annexures, and a prayer. Contains no assessment of
                           the firm's own confidence, no exposure arithmetic,
                           no unverified citation, and no trace of how it was
                           prepared.

    build_file_note()      stays in the office. Contains everything the other
                           document must not: postures and why, weaknesses,
                           exposure, evidence gaps, unverified authorities,
                           panel disagreements, and the board summary.

STRUCTURE OF THE FILING DOCUMENT
--------------------------------
The A-to-O framework used before Adjudicating Authorities, Commissioner
(Appeals) and the GSTAT — cause title, Disputes at a Glance, issue-wise reply,
consolidated payments, evidentiary index, prayer. Not a memo with a draft in
it. A document that can be signed and uploaded.

Typography is deliberately plain: Arial 11pt, black on white, bold for
headings and nothing else. A document that looks designed looks amateur.
"""

import io
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from . import config, defects
from .verification import (
    ACTIONABLE,
    NOT_FOUND,
    SUPERSEDED,
    UNVERIFIED,
    VERIFIED,
    extract_citations,
)

BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)

BODY_FONT = "Arial"
BODY_SIZE = Pt(11)

STATUS_LABEL = {
    VERIFIED: "Verified",
    SUPERSEDED: "Superseded",
    UNVERIFIED: "To be confirmed",
    NOT_FOUND: "Not traced",
}

STRENGTH_LABEL = {
    "strong": "Strong",
    "defensible": "Defensible",
    "weak": "Weak",
}

CONFIDENCE_LABEL = {
    "strong": "Strong",
    "defensible": "Defensible",
    "weak": "Weak",
    "insufficient_information": "Further information required",
}

ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI",
         "XII", "XIII", "XIV", "XV", "XVI", "XVII", "XVIII", "XIX", "XX"]


def _roman(n: int) -> str:
    return ROMAN[n - 1] if 1 <= n <= len(ROMAN) else str(n)


# ---------------------------------------------------------------------------
# Typography
# ---------------------------------------------------------------------------


def _configure_styles(doc: Document):
    """Arial 11pt, black, throughout. Headings differ only by weight."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = BLACK
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name in ("Heading 1", "Heading 2", "Heading 3", "Title"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = BODY_FONT
        style.font.color.rgb = BLACK
        style.font.bold = True
        style.font.italic = False
        style.font.size = Pt(13) if name in ("Heading 1", "Title") else Pt(11)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(4)

    for name in ("List Bullet", "List Number"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE
        style.font.color.rgb = BLACK


def _para(doc: Document, text: str = "", bold: bool = False,
          italic: bool = False, size: Pt = None, align=None, space_after=None,
          indent=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    if space_after is not None:
        paragraph.paragraph_format.space_after = space_after
    if indent is not None:
        paragraph.paragraph_format.left_indent = indent
    if text:
        run = paragraph.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = size or BODY_SIZE
        run.font.color.rgb = BLACK
        run.bold = bold
        run.italic = italic
    return paragraph


def _heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper() if level == 1 else text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    run.bold = True
    return paragraph


def _rule(doc: Document):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)
    return paragraph


def _plain_table(doc: Document, columns: int):
    table = doc.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    return table


def _set_cell(cell, text: str, bold: bool = False, size: Pt = None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(str(text) if text not in (None, "") else "—")
    run.font.name = BODY_FONT
    run.font.size = size or Pt(10)
    run.font.color.rgb = BLACK
    run.bold = bold


def _grid(doc: Document, headers: List[str], rows: List[List[str]],
          widths: List[float] = None, bold_columns: Tuple[int, ...] = ()):
    """A bordered table with a bold header row — how a reply states figures."""
    if not rows:
        return None
    table = _plain_table(doc, len(headers))
    if widths:
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)
    header_cells = table.add_row().cells
    for cell, label in zip(header_cells, headers):
        _set_cell(cell, label, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row[:len(headers)]):
            _set_cell(cells[index], value, bold=index in bold_columns)
    _para(doc, space_after=Pt(4))
    return table


def _particulars(doc: Document, rows):
    rows = [(k, v) for k, v in rows if v not in (None, "", [])]
    if not rows:
        return
    table = _plain_table(doc, 2)
    table.columns[0].width = Inches(2.1)
    table.columns[1].width = Inches(4.3)
    for key, value in rows:
        cells = table.add_row().cells
        _set_cell(cells[0], key, bold=True)
        _set_cell(cells[1], value)
    _para(doc, space_after=Pt(4))


def _numbered_body(doc: Document, text: str, indent_continuations: bool = True):
    """Render pre-numbered text, indenting continuation lines."""
    for block in str(text or "").split("\n"):
        stripped = block.strip()
        if not stripped:
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        numbered = stripped[:1].isdigit() and (
            "." in stripped[:6] or ")" in stripped[:6]
        )
        if indent_continuations and numbered:
            paragraph.paragraph_format.left_indent = Inches(0.35)
            paragraph.paragraph_format.first_line_indent = Inches(-0.35)
        run = paragraph.add_run(stripped)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
        run.font.color.rgb = BLACK


def _page_footer(doc: Document, left: str):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{left}    |    Page ")
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def _page_header(doc: Document, text: str):
    """A running header — used to stamp a draft as not for filing."""
    if not text:
        return
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK
    run.bold = True


def _margins(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


# ---------------------------------------------------------------------------
# Shared reading of a matter
# ---------------------------------------------------------------------------


def _unpack(matter: Dict[str, Any]):
    intake = matter.get("intake", {})
    result = matter.get("result") or {}
    determination = result.get("determination") or {}
    verification = result.get("verification") or {}
    metadata = matter.get("metadata") or {}
    defect_list = determination.get("defects") or intake.get("defects") or []
    # Both builders index every limb as a mapping. The panel already drops
    # malformed entries, but a stored matter is read back from disk and may
    # predate that, and neither document may fail to build over one bad limb.
    defect_list = [d for d in defect_list if isinstance(d, dict)]
    return intake, determination, verification, metadata, defect_list


def _verified_index(verification: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    """Citation text -> its verification record, for gating the filing document."""
    index = {}
    for authority in verification.get("authorities") or []:
        citation = (authority.get("citation") or "").strip().lower()
        if citation:
            index[citation] = authority
    return index


def _is_filable(citation: str, index: Dict[str, Dict[str, Any]]) -> bool:
    """
    May this authority appear in the document that goes to the officer?

    An authority nobody checked is not filable. If verification did not run at
    all the index is empty and NOTHING is filable, which is the correct
    failure direction: a reply with fewer authorities is recoverable, a reply
    with a fabricated one is not.
    """
    record = index.get((citation or "").strip().lower())
    return bool(record and record.get("status") == VERIFIED)


def _reply_form(intake: Dict[str, Any], pack_types: Dict[str, Any]) -> str:
    notice = pack_types.get(intake.get("notice_type"))
    return (notice.reply_form if notice else None) or "the appropriate reply form"


def _prose_citation_gaps(
    determination: Dict[str, Any],
    defect_list: List[Dict[str, Any]],
    verified: Dict[str, Dict[str, Any]],
    pack,
) -> List[Dict[str, str]]:
    """
    Citations sitting INSIDE the filed prose that are not filable.

    The structured lists are gated by withholding: an authority that failed
    verification simply does not print. A citation the chairman wrote into a
    submission paragraph cannot be withheld without rewriting the paragraph,
    so it is surfaced instead — as a blocker in the file note and as a stamp
    on the filing document — until a reviewer confirms it or strikes it.

    Same failure direction as `_is_filable`: a citation nobody checked is a
    gap, so a run where verification never happened stamps rather than files.
    """
    texts: List[Tuple[str, str]] = []
    for defect in defect_list:
        label = (f"defect {defect.get('index')} "
                 f"({defect.get('heading') or 'unheaded'})")
        for field in ("submission", "facts", "our_position"):
            if defect.get(field):
                texts.append((label, str(defect[field])))
    if determination.get("preliminary_submissions"):
        texts.append(("the preliminary submissions",
                      str(determination["preliminary_submissions"])))

    gaps: List[Dict[str, str]] = []
    seen = set()
    for label, text in texts:
        for citation in extract_citations(text, pack):
            if _is_filable(citation, verified):
                continue
            key = re.sub(r"\s+", " ", citation).lower()
            if key in seen:
                continue
            seen.add(key)
            record = verified.get(citation.strip().lower()) or {}
            status = STATUS_LABEL.get(record.get("status"),
                                      "not covered by verification")
            gaps.append({
                "citation": citation,
                "message": (
                    f"The filed text for {label} cites \"{citation}\" "
                    f"({status}). It sits inside a filed paragraph and cannot "
                    "be withheld automatically — confirm it against the "
                    "reported text or strike it from the prose before filing."
                ),
            })
    return gaps


PROSE_GAP_STAMP = ("NOT FOR FILING — UNVERIFIED AUTHORITY CITED IN THE TEXT "
                   "— SEE THE FILE NOTE")


# ---------------------------------------------------------------------------
# 1. The filing document
# ---------------------------------------------------------------------------


def build_filing_reply(matter: Dict[str, Any]) -> bytes:
    """
    The document that is signed and filed.

    Everything in here is written to be read by the proper officer. Nothing in
    here assesses the firm's own confidence, quantifies the client's exposure,
    or carries an authority that failed verification.
    """
    from .domains import get_pack
    intake, determination, verification, metadata, defect_list = _unpack(matter)
    pack = get_pack(matter.get("domain", "gst"))
    verified = _verified_index(verification)

    doc = Document()
    _configure_styles(doc)
    _margins(doc)

    watermark = metadata.get("watermark")
    # A citation inside the filed prose that did not verify cannot be
    # withheld the way a table entry is. Blocking the export entirely only
    # pushes the text out through the clipboard, so the document exports —
    # stamped on every page so it cannot be mistaken for filable.
    if _prose_citation_gaps(determination, defect_list, verified, pack):
        watermark = (f"{watermark} — {PROSE_GAP_STAMP}" if watermark
                     else PROSE_GAP_STAMP)
    if watermark:
        _page_header(doc, watermark)

    client = intake.get("client_name") or "the Noticee"
    reference = intake.get("notice_reference") or str(matter.get("id", ""))[:8].upper()
    _page_footer(doc, f"{client} — {intake.get('tax_period', '')}".strip(" —"))

    # ---- Cause title -----------------------------------------------------
    forum = _forum_line(intake)
    for line in forum:
        _para(doc, line, bold=True, size=Pt(12),
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    _para(doc, space_after=Pt(4))
    _rule(doc)

    _particulars(doc, [
        ("Noticee", f"Tvl. {client}" if intake.get("state") == "Tamil Nadu"
                    else f"M/s. {client}"),
        ("GSTIN", intake.get("gstin")),
        ("Period in dispute", intake.get("tax_period")),
        ("Impugned notice", _notice_line(intake, pack)),
        ("Reference", reference),
        ("ARN", intake.get("notice_arn")),
        ("DIN", intake.get("din")),
        ("Amount in dispute", _rupees(_matter_amount(defect_list, intake))),
        ("Reply due", _date(intake.get("due_date"))),
        ("Filed in", _reply_form(intake, pack.NOTICE_TYPES)),
    ])

    _para(doc, "COMPREHENSIVE WRITTEN REPLY", bold=True, size=Pt(12),
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _rule(doc)

    # ---- A. Preliminary submissions --------------------------------------
    _heading(doc, "A.  Preliminary Submissions", 1)
    preliminary = determination.get("preliminary_submissions")
    if preliminary:
        _numbered_body(doc, preliminary)
    else:
        _para(doc, _default_preliminary(intake, client, pack))

    # ---- B. Disputes at a glance -----------------------------------------
    if defect_list:
        _heading(doc, "B.  Disputes at a Glance", 1)
        rows = []
        for defect in defect_list:
            amount = defects.defect_total(defect)
            allegation = defect.get("heading") or ""
            if amount:
                allegation += f" — Rs. {defects.indian_number(amount)}"
            rows.append([
                str(defect.get("index", "")),
                allegation,
                defect.get("our_position") or defects.POSTURE_LABEL.get(
                    defect.get("posture"), ""),
            ])
        _grid(doc, ["Sl.", "Allegation in the notice", "Noticee's position"],
              rows, widths=[0.4, 2.9, 3.1], bold_columns=(2,))

    # ---- C. Issue-wise reply ---------------------------------------------
    if defect_list:
        _heading(doc, "C.  Issue-wise Detailed Reply", 1)
        for defect in defect_list:
            _defect_section(doc, defect, verified)

    # ---- D. Consolidated statement of payments ---------------------------
    payment_rows = _payment_rows(defect_list)
    if payment_rows:
        _heading(doc, "D.  Consolidated Statement of Payments", 1)
        _para(doc, "The following amounts have been discharged through FORM "
                   "GST DRC-03:")
        _grid(doc,
              ["Sl.", "Issue", "Amount (Rs.)", "Reference", "Remarks"],
              payment_rows, widths=[0.4, 2.2, 1.2, 1.6, 1.0])
        if any("protest" in (row[4] or "").lower() for row in payment_rows):
            _para(doc,
                  "The payments marked as made under protest are made without "
                  "prejudice to the Noticee's position that the said input tax "
                  "credit is eligible, and do not constitute any admission of "
                  "liability. The Noticee expressly reserves its right to seek "
                  "refund of the amounts so paid.")

    # ---- E. Evidentiary index --------------------------------------------
    annexures = _annexure_rows(defect_list)
    if annexures:
        _heading(doc, "E.  Evidentiary Index — Documents Enclosed", 1)
        _grid(doc, ["Annexure", "Document", "Issue supported"], annexures,
              widths=[0.9, 3.6, 1.9])

    # ---- F. Prayer --------------------------------------------------------
    _heading(doc, "F.  Prayer", 1)
    _para(doc,
          "In view of the foregoing submissions, the Noticee most respectfully "
          "prays that the Honourable Proper Officer may be pleased to:")
    prayer_rows = []
    counter = 0
    for defect in defect_list:
        relief = defect.get("prayer_relief")
        if not relief:
            relief = _default_relief(defect)
        counter += 1
        prayer_rows.append([_roman(counter), relief])
    counter += 1
    prayer_rows.append([
        _roman(counter),
        "GRANT a personal hearing under Section 75(4) of the Central Goods and "
        "Services Tax Act, 2017 at a mutually convenient date before any "
        "adverse view is taken.",
    ])
    counter += 1
    prayer_rows.append([
        _roman(counter),
        "PASS such further order or grant such other relief as this Honourable "
        "Authority may deem fit and proper in the facts and circumstances of "
        "the case.",
    ])
    _grid(doc, ["Sl.", "Relief prayed"], prayer_rows, widths=[0.6, 5.8])

    _para(doc, "And for this act of justice, the Noticee as in duty bound "
               "shall ever pray.", space_after=Pt(18))

    # ---- Signature block --------------------------------------------------
    #
    # The chairman prompt used to forbid this block, on the reasoning that "the
    # firm supplies those". The firm should not have to: a reply without an
    # addressee, a place, a date and a signatory is not a reply, and every
    # exported document arrived needing the same manual surgery.
    _para(doc, "Yours faithfully,", space_after=Pt(2))
    _para(doc, f"For {'Tvl.' if intake.get('state') == 'Tamil Nadu' else 'M/s.'} "
               f"{client}", bold=True, space_after=Pt(30))
    _para(doc, "Authorised Signatory", space_after=Pt(2))
    _particulars(doc, [
        ("Place", intake.get("place") or ""),
        ("Date", ""),
    ])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _forum_line(intake: Dict[str, Any]) -> List[str]:
    office = intake.get("jurisdiction_office") or ""
    officer = intake.get("issuing_officer") or ""
    designation = ""
    for candidate in ("Assistant Commissioner", "Deputy Commissioner",
                      "Joint Commissioner", "Additional Commissioner",
                      "Superintendent", "Commissioner"):
        if candidate.lower() in officer.lower():
            designation = candidate
            break
    state = intake.get("state") or ""
    lines = [f"BEFORE THE {(designation or 'PROPER OFFICER').upper()}"
             f"{' (ST)' if state and designation else ''}"]
    if office:
        lines.append(office.upper())
    if state:
        lines.append(f"COMMERCIAL TAXES DEPARTMENT, GOVERNMENT OF {state.upper()}")
    return lines


def _notice_line(intake: Dict[str, Any], pack) -> str:
    notice = pack.NOTICE_TYPES.get(intake.get("notice_type"))
    parts = []
    if notice:
        parts.append(f"FORM GST {notice.code} — {notice.name}")
        parts.append(f"issued under {notice.statute}")
    if intake.get("notice_date"):
        parts.append(f"dated {_date(intake['notice_date'])}")
    return ", ".join(parts)


def _default_preliminary(intake: Dict[str, Any], client: str, pack) -> str:
    notice = pack.NOTICE_TYPES.get(intake.get("notice_type"))
    form = notice.code if notice else "the notice"
    statute = notice.statute if notice else "the provisions cited therein"
    return (
        f"1.  M/s. {client}, holding GSTIN {intake.get('gstin', '')} "
        f"(hereinafter referred to as 'the Noticee'), registered under the "
        f"Central Goods and Services Tax Act, 2017 and the corresponding State "
        f"Act, hereby submits this comprehensive written reply to the notice "
        f"issued in FORM GST {form} under {statute}, in respect of the tax "
        f"period {intake.get('tax_period', '')}.\n\n"
        "2.  The Noticee has examined each discrepancy raised in the notice and "
        "submits its response issue-wise hereunder. The Noticee has cooperated "
        "fully with the proceedings and undertakes to produce such further "
        "documentary evidence as may be called for."
    )


def _defect_section(doc: Document, defect: Dict[str, Any],
                    verified: Dict[str, Dict[str, Any]]):
    index = defect.get("index", "")
    _heading(doc, f"Issue {_roman(int(index)) if str(index).isdigit() else index}: "
                  f"{defect.get('heading', '')}", 2)

    amount = defects.defect_total(defect)
    if amount:
        heads = defects.normalise_heads(defect.get("amount_by_head"))
        allocated = any(heads.get(h) for h in defects.TAX_HEADS)
        split = (f" ({defects.format_heads(heads)})" if allocated else "")
        _para(doc,
              f"Amount in issue: Rs. {defects.indian_number(amount)}{split}.",
              bold=True, size=Pt(10))

    contention = _clean_contention(defect.get("department_contention"))
    if contention:
        _para(doc, "Department's allegation", bold=True, size=Pt(10))
        _para(doc, contention)

    if defect.get("facts"):
        _para(doc, "Factual position", bold=True, size=Pt(10))
        _numbered_body(doc, defect["facts"])

    # Gated exactly as the authorities table below is: the chairman is invited
    # to put circulars, notifications and case law in this field, and an entry
    # that did not come back VERIFIED belongs in the file note, not here. If
    # verification never ran, nothing is filable — the correct direction.
    framework = [
        entry for entry in (defect.get("legal_framework") or [])
        if isinstance(entry, dict) and entry.get("provision")
        and _is_filable(entry.get("provision", ""), verified)
    ]
    if framework:
        _para(doc, "Legal framework", bold=True, size=Pt(10))
        _grid(doc, ["Provision", "Relevance"],
              [[e.get("provision", ""), e.get("relevance", "")] for e in framework],
              widths=[2.4, 4.0])

    # Only VERIFIED authority reaches this document. Anything else is in the
    # file note with a confirm-before-filing flag, which is where a caveat
    # belongs — not in front of the officer.
    filable = [
        a for a in (defect.get("authorities") or [])
        if isinstance(a, dict) and _is_filable(a.get("citation", ""), verified)
    ]
    if filable:
        _para(doc, "Authorities relied upon", bold=True, size=Pt(10))
        _grid(doc, ["Authority", "Proposition"],
              [[a.get("citation", ""), a.get("proposition", "")] for a in filable],
              widths=[2.8, 3.6])

    splits = [s for s in (defect.get("splits") or []) if isinstance(s, dict)]
    if splits:
        _para(doc, "Position taken, item by item", bold=True, size=Pt(10))
        _grid(doc, ["Particulars", "Amount (Rs.)", "Position taken"],
              [[s.get("description", ""),
                defects.indian_number(defects.total_of(s.get("amount_by_head"))),
                defects.POSTURE_LABEL.get(s.get("posture"), s.get("posture", ""))]
               for s in splits],
              widths=[3.0, 1.4, 2.0])

    submission = defect.get("submission")
    if submission:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(12)
        run = paragraph.add_run("SUBMISSION: ")
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
        run.bold = True
        run.font.color.rgb = BLACK
        body = paragraph.add_run(submission)
        body.font.name = BODY_FONT
        body.font.size = BODY_SIZE
        body.bold = True
        body.font.color.rgb = BLACK


def _clean_contention(text: str) -> str:
    """
    Render the department's allegation as prose, never as a flattened table.

    When the chairman supplies a contention it is already a sentence or two.
    When it falls back to the raw notice extract — because the chairman omitted
    it, or the run degraded — that extract can be an entire annexure with its
    column rulers, and pasting it into a filed reply as a paragraph of loose
    digits reads as though nobody looked at the document.

    So sentences that are mostly numbers are dropped. What survives is the
    department's own words; the figures are already stated, correctly and
    head-wise, in the line above.
    """
    raw = re.sub(r"\s+", " ", str(text or "")).strip().lstrip("•‣▪ ")
    if not raw:
        return ""

    kept = []
    for sentence in re.split(r"(?<=[.:])\s+", raw):
        sentence = sentence.strip()
        if len(sentence) < 3:
            continue
        if _is_table_residue(sentence):
            continue
        kept.append(sentence)
        if sum(len(s) for s in kept) > 700:
            break

    return " ".join(kept)


# Four or more standalone numbers in a row. This is the signature of a table
# flattened into text — a column ruler, or a row of head-wise figures — and it
# does not occur in prose. A digit-ratio threshold was tried first and let a
# whole annexure through, because the long column descriptions the department
# writes ("ITC on inward supplies other than imports and inward supplies liable
# to reverse charge…") dilute the ratio below any workable cut-off.
_NUMBER_RUN_RE = re.compile(r"(?:(?<![\w.])[\d,]+(?:\.\d+)?(?![\w.])\s+){3,}"
                            r"(?<![\w.])[\d,]+(?:\.\d+)?(?![\w.])")


# Column headings the department prints above every annexure. They survive the
# numeric filters because they carry no digits at all, and they read as
# gibberish in a filed paragraph.
_COLUMN_HEADER_RE = re.compile(
    r"\b(?:S\.?\s?No\.?|Sl\.?\s?No\.?|Table\s+No\.?|Return\s*/\s*Statement"
    r"|Taxable\s+value\s+SGST|SGST\s+CGST\s+IGST)\b",
    re.IGNORECASE,
)


def _is_table_residue(sentence: str) -> bool:
    if _NUMBER_RUN_RE.search(sentence):
        return True
    if _COLUMN_HEADER_RE.search(sentence):
        return True
    digits = sum(c.isdigit() for c in sentence)
    return digits > len(sentence) * 0.30


def _payment_rows(defect_list: List[Dict[str, Any]]) -> List[List[str]]:
    rows = []
    counter = 0
    for defect in defect_list:
        payment = defect.get("payment")
        # Guarded like splits, authorities and legal_framework above: a stored
        # matter can carry a payment the panel wrote as prose, and this
        # document must still build.
        if not isinstance(payment, dict) or not payment.get("reference"):
            continue
        counter += 1
        under_protest = bool(payment.get("under_protest")) or \
            defect.get("posture") == defects.PAID_UNDER_PROTEST
        amount = defects.total_of(payment.get("tax_by_head")) or \
            defects.defect_total(defect)
        reference = payment.get("reference", "")
        if payment.get("date"):
            reference = f"{reference} dated {payment['date']}"
        rows.append([
            str(counter),
            f"Issue {defect.get('index')} — {defect.get('heading', '')}",
            defects.indian_number(amount),
            reference,
            "Paid under protest" if under_protest else "Agreed and paid",
        ])
    return rows


def _annexure_rows(defect_list: List[Dict[str, Any]]) -> List[List[str]]:
    rows = []
    counter = 0
    for defect in defect_list:
        for annexure in defect.get("annexures") or []:
            counter += 1
            rows.append([
                f"Annexure-{counter}",
                str(annexure),
                f"Issue {defect.get('index')} — {defect.get('heading', '')}",
            ])
    return rows


def _default_relief(defect: Dict[str, Any]) -> str:
    verb = defects.POSTURE_RELIEF_VERB.get(defect.get("posture"), "CONSIDER")
    amount = defects.defect_total(defect)
    tail = f" of Rs. {defects.indian_number(amount)}" if amount else ""
    return (f"{verb} the allegation at Issue {defect.get('index')} — "
            f"{defect.get('heading', '')}{tail}.")


UNREAD_LABEL = "Not read from the notice"


def _defect_amount_cell(defect: Dict[str, Any]) -> str:
    """
    A limb's amount for a table, or an honest blank.

    A limb whose figure could not be read carries `amount_unread`, and every
    render site printed `defects.defect_total()` regardless — which returns 0,
    so an unread figure appeared as "Rs. 0" or "0" in the defect register and
    the hearing brief. A zero is a figure; the reviewer reads it as "nothing in
    issue on this limb" and moves on, which is the one conclusion that is
    certainly wrong. `amount_note` — which the chairman is instructed to fill
    for exactly this case — was rendered nowhere at all.
    """
    if defect.get("amount_unread"):
        note = (defect.get("amount_note") or "").strip()
        return f"{UNREAD_LABEL} — {note}" if note else UNREAD_LABEL
    return defects.indian_number(defects.defect_total(defect))


def _matter_amount_is_partial(defect_list) -> bool:
    """True when any limb's figure is unread, so the limb sum understates."""
    return any(d.get("amount_unread") for d in (defect_list or [])
               if isinstance(d, dict))


def _file_note_amount(defect_list, intake) -> Optional[str]:
    """
    The matter total for the file note, marked incomplete where it is.

    The filing document quotes the department's own total; the reviewer needs
    to know when our arithmetic could not reach it, because the gap is the
    measure of what still has to be taken off the annexure.
    """
    shown = _rupees(_matter_amount(defect_list, intake))
    if not _matter_amount_is_partial(defect_list):
        return shown
    unread = sum(1 for d in defect_list
                 if isinstance(d, dict) and d.get("amount_unread"))
    marker = (f"INCOMPLETE — {unread} limb(s) carry no figure that could be "
              "read from the notice")
    return f"{shown} ({marker})" if shown else marker


def _matter_amount(defect_list, intake) -> Optional[float]:
    """
    The amount in dispute across the matter.

    Where a limb's figure could not be read, the sum of the limbs is an
    UNDERSTATEMENT — `defect_total` returns 0 for an unread limb. Quoting that
    understated figure back to the department as "Amount in dispute" states a
    number that is wrong and in the taxpayer's favour, which is exactly the
    kind of error that costs credibility on the limbs that matter. So when any
    limb is unread the notice's own declared total is preferred: it is the
    department's figure, not ours.
    """
    if defect_list:
        if _matter_amount_is_partial(defect_list) and \
                intake.get("amount_disputed"):
            return intake["amount_disputed"]
        return defects.matter_total(defect_list)
    return intake.get("amount_disputed")


# ---------------------------------------------------------------------------
# 2. The internal file note
# ---------------------------------------------------------------------------


def build_file_note(matter: Dict[str, Any]) -> bytes:
    """
    The working paper. Never leaves the firm.

    This is where every hedge, every exposure figure, every unverified citation
    and every evidence gap belongs — stated bluntly, because the only reader is
    a professional deciding whether to sign.
    """
    from .domains import get_pack
    intake, determination, verification, metadata, defect_list = _unpack(matter)
    pack = get_pack(matter.get("domain", "gst"))
    verified = _verified_index(verification)

    doc = Document()
    _configure_styles(doc)
    _margins(doc)
    _page_header(doc, "INTERNAL WORKING PAPER — NOT FOR SUBMISSION")
    _page_footer(doc, f"Ref: {str(matter.get('id', ''))[:8].upper()}")

    if config.FIRM_NAME:
        _para(doc, config.FIRM_NAME, bold=True, size=Pt(13),
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
        _para(doc, config.FIRM_SUBTITLE, size=Pt(10),
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
        _rule(doc)

    _para(doc, "FILE NOTE — NOTICE REPLY", bold=True, size=Pt(13),
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    _para(doc,
          "Internal working paper. This document must not be sent to the "
          "client, annexed to the reply, or filed with the department. It "
          "records the firm's own assessment, including weaknesses, exposure "
          "and unresolved points.",
          italic=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _rule(doc)

    _particulars(doc, [
        ("Client", intake.get("client_name")),
        ("GSTIN", intake.get("gstin")),
        ("Notice", intake.get("notice_type")),
        ("Reference", intake.get("notice_reference")),
        ("Jurisdiction", intake.get("state")),
        ("Tax period", intake.get("tax_period")),
        ("Amount in dispute", _file_note_amount(defect_list, intake)),
        ("Reply due", _date(intake.get("due_date"))),
        ("Prepared on", datetime.now(timezone.utc).strftime("%d %B %Y")),
        ("Matter reference", str(matter.get("id", ""))[:8].upper()),
    ])

    _para(doc, config.EXPORT_REVIEW_NOTE, italic=True, size=Pt(9))

    # ---- 1. Blockers ------------------------------------------------------
    blockers = list(determination.get("filing_blockers") or [])
    # Recomputed here as well as at panel time, so a stored matter that
    # predates the panel-side check still surfaces a prose citation that was
    # never verified. Gaps already recorded by the panel (matched on the
    # citation text) are not listed twice.
    recorded = "\n".join(str(b) for b in blockers)
    for gap in _prose_citation_gaps(determination, defect_list, verified, pack):
        if gap["citation"] not in recorded:
            blockers.append(gap["message"])

    # An unread figure is a blocker in its own right. It reached the reviewer
    # as "Rs. 0" in the defect register before this, which reads as "nothing in
    # issue on this limb" — the one conclusion that is certainly wrong.
    for defect in defect_list:
        if not defect.get("amount_unread"):
            continue
        note = (defect.get("amount_note") or "").strip()
        blockers.append(
            f"Limb {defect.get('index')} ({defect.get('heading') or 'unheaded'}"
            "): the amount could not be read from the notice. Take it from the "
            "department's annexure and enter it before filing — the reply "
            "currently answers this limb without stating a figure."
            + (f" Panel note: {note}" if note else "")
        )
    _heading(doc, "1.  Before this reply can be filed", 1)
    if blockers:
        _para(doc, f"{len(blockers)} matter(s) must be resolved before filing.",
              bold=True)
        for index, blocker in enumerate(blockers, start=1):
            _para(doc, f"{index}.  {blocker}", indent=Inches(0.3))
    else:
        _para(doc, "No structural blockers were identified. The substantive "
                   "points below still require partner judgement.")

    # ---- 2. Position and triage -------------------------------------------
    _heading(doc, "2.  Position Recommended", 1)
    _para(doc, determination.get("recommended_position") or
          "No recommendation was settled on the material available.")

    triage = determination.get("triage") or defects.triage(defect_list)
    rows = []
    if determination.get("confidence"):
        rows.append(("Assessment", CONFIDENCE_LABEL.get(
            determination["confidence"], determination["confidence"].title())))
    if determination.get("lead_argument"):
        rows.append(("Argument taken first", determination["lead_argument"]))
    if triage.get("total_count"):
        rows.append((
            "Limbs",
            f"{triage['total_count']} defects, Rs. "
            f"{defects.indian_number(triage['total_amount'])}. "
            f"{triage['argue_count']} argued (Rs. "
            f"{defects.indian_number(triage['argued_amount'])}), "
            f"{triage['settle_count']} settled on documents or payment (Rs. "
            f"{defects.indian_number(triage['settled_amount'])})."
        ))
    _particulars(doc, rows)

    # ---- 3. Defect register ------------------------------------------------
    if defect_list:
        _heading(doc, "3.  Defect Register", 1)
        _grid(doc,
              ["Sl.", "Defect", "Amount (Rs.)", "Posture", "Strength"],
              [[str(d.get("index", "")),
                d.get("heading", ""),
                _defect_amount_cell(d),
                defects.POSTURE_LABEL.get(d.get("posture"), d.get("posture", "")),
                STRENGTH_LABEL.get((d.get("strength") or "").lower(),
                                   d.get("strength") or "—")]
               for d in defect_list],
              widths=[0.4, 2.5, 1.1, 1.4, 1.0])

    # ---- 4. Evidence gaps — the most important section --------------------
    gaps = [(d, list(d.get("evidence_gap") or [])) for d in defect_list]
    gaps = [(d, g) for d, g in gaps if g]
    _heading(doc, "4.  Evidence Gaps", 1)
    _para(doc,
          "Documents the officer will require which the engagement team has "
          "not confirmed are available. A limb with a sound argument and a "
          "missing document is a limb that will be lost — this is the section "
          "to clear before the reply is filed, not after.",
          italic=True, size=Pt(9))
    if gaps:
        _grid(doc, ["Defect", "Document not confirmed as held"],
              [[f"{d.get('index')} — {d.get('heading', '')}", item]
               for d, items in gaps for item in items],
              widths=[2.4, 4.0])
    else:
        _para(doc, "No evidence gap was recorded. Confirm this positively "
                   "against the evidence list for each defect rather than "
                   "reading silence as sufficiency.")

    # ---- 5. Schedule of authorities ---------------------------------------
    authorities = verification.get("authorities") or []
    _heading(doc, "5.  Schedule of Authorities", 1)
    if authorities:
        summary = verification.get("summary") or {}
        outstanding = sum(summary.get(k, 0) or 0
                          for k in ("superseded", "unverified", "not_found"))
        if outstanding:
            _para(doc,
                  f"{outstanding} of {summary.get('total', len(authorities))} "
                  "authorities did not verify. These have been WITHHELD from "
                  "the filing document. Confirm them against the reported text "
                  "if any is to be relied upon.", bold=True)
        _grid(doc, ["Authority", "Cited for", "Defect", "Status", "Remarks"],
              [[a.get("citation", ""),
                a.get("proposition", ""),
                str(a.get("defect_index") or "—"),
                STATUS_LABEL.get(a.get("status", UNVERIFIED), "To be confirmed"),
                (a.get("note", "") +
                 (f" Read as: {a['correction']}" if a.get("correction") else ""))]
               for a in authorities],
              widths=[1.8, 1.7, 0.5, 1.0, 1.4])
        _para(doc,
              "Only authorities shown as 'Verified' appear in the filing "
              "document. 'Superseded', 'To be confirmed' and 'Not traced' are "
              "withheld from it by design.",
              italic=True, size=Pt(9))
    else:
        _para(doc,
              "No authority was cited, or verification did not run. The filing "
              "document therefore carries no case law. Consider whether the "
              "contested limbs can be sustained on the statute alone.")

    # ---- 6. Risk flags and open questions ---------------------------------
    risk_flags = determination.get("risk_flags") or []
    open_questions = determination.get("open_questions") or []
    if risk_flags or open_questions:
        _heading(doc, "6.  Points for Partner Attention", 1)
        counter = 0
        for flag in risk_flags:
            counter += 1
            _para(doc, f"{counter}.  {flag}", indent=Inches(0.3))
        for question in open_questions:
            counter += 1
            _para(doc, f"{counter}.  {question}", indent=Inches(0.3))

    # ---- 7. Documents to obtain -------------------------------------------
    documents = determination.get("documents_to_collect") or []
    if documents:
        _heading(doc, "7.  Documents to be Obtained", 1)
        for document in documents:
            _para(doc, f"•  {document}", indent=Inches(0.2))

    # ---- 8. Working note and disagreements --------------------------------
    if determination.get("working_note") or determination.get("panel_disagreements"):
        doc.add_page_break()
        _heading(doc, "8.  Basis on which the Position was Settled", 1)
        _para(doc,
              "The peer-review and litigation-defence record: what was argued, "
              "what was rejected, and on what reasoning.",
              italic=True, size=Pt(9))
        if determination.get("working_note"):
            _numbered_body(doc, determination["working_note"])

        alternatives = [e for e in (determination.get("panel_disagreements") or [])
                        if isinstance(e, dict)]
        if alternatives:
            _heading(doc, "Positions considered and not adopted", 2)
            for entry in alternatives:
                _particulars(doc, [
                    ("Question", entry.get("question")),
                    ("Positions considered", entry.get("positions")),
                    ("Basis on which settled", entry.get("resolution")),
                ])

    # ---- 9. Statutory computations -----------------------------------------
    _computations_section(doc, matter)

    # ---- 10. Hearing brief -------------------------------------------------
    _hearing_brief_section(doc, matter, defect_list)

    # ---- 11. Board summary -------------------------------------------------
    if determination.get("board_summary"):
        _heading(doc, "11.  Summary for the Board", 1)
        _para(doc, determination["board_summary"])

    if determination.get("unstructured_output"):
        doc.add_page_break()
        _heading(doc, "Appendix — Unstructured Chairman Output", 1)
        _para(doc, "The determination could not be parsed into the structured "
                   "shape. The raw output is reproduced so nothing is lost.",
              italic=True, size=Pt(9))
        _numbered_body(doc, determination["unstructured_output"],
                       indent_continuations=False)

    if config.EXPORT_PROVENANCE:
        doc.add_page_break()
        _heading(doc, "Appendix — Panel Composition", 1)
        models = metadata.get("models") or {}
        _particulars(doc, [(k.title(), v) for k, v in models.items()] + [
            ("Client particulars withheld in preparation",
             "Yes" if metadata.get("anonymised") else "No"),
            ("Matter reference", matter.get("id")),
        ])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# File-note sections that are computed rather than argued
# ---------------------------------------------------------------------------
# Both of these belong to the INTERNAL document and to no other. Interest and
# penalty stages are the firm's own workings and quantify what the matter costs
# if it goes badly; putting a penalty computation in front of the officer
# volunteers an admission nobody asked for. The hearing brief is a script for
# the firm's own representative.


def _computations_section(doc: Document, matter: Dict[str, Any]):
    """Interest, penalty stages, pre-deposit and amnesty — with their working."""
    from . import calculators

    try:
        computed = calculators.matter_computations(matter)
    except Exception:
        # A working note is worth having without this section; it is not worth
        # losing the whole document to an arithmetic edge case.
        return

    blocks = computed.get("computations") or {}
    penalty = blocks.get("penalty") or {}
    amnesty = blocks.get("amnesty_128a") or {}
    limitation = blocks.get("appeal_limitation") or {}
    deposit = blocks.get("predeposit_107") or {}

    if not any([penalty.get("computed"), amnesty.get("reasons"),
                limitation.get("computed"), deposit.get("computed")]):
        return

    _heading(doc, "9.  Statutory Computations", 1)
    _para(doc,
          "Computed locally under the provisions named — not by a model. Each "
          "figure carries its own working and is to be checked against the "
          "electronic liability register before it is offered to the "
          "department.",
          italic=True, size=Pt(9))

    if penalty.get("computed"):
        _heading(doc, f"Penalty exposure — Section {penalty['section']}", 2)
        _para(doc, f"Computed on tax of {_rupees(penalty.get('tax'))}.",
              size=Pt(9), italic=True)
        rows = []
        for stage in penalty.get("stages") or []:
            when = {
                "before_notice": "If paid before the notice",
                "within_30_days": "If paid within 30 days of the notice",
                "on_order": "On determination by order",
            }.get(stage["stage"], stage["stage"])
            if stage.get("deadline"):
                when += f" (by {_date(stage['deadline'])})"
            rows.append([when, f"{stage['rate']:g}%",
                         _rupees(stage["amount"]) or "Nil"])
        _grid(doc, ["Stage", "Rate", "Penalty"], rows,
              widths=[3.4, 0.9, 1.8], bold_columns=(2,))
        for caveat in penalty.get("caveats") or []:
            _para(doc, f"•  {caveat}", size=Pt(9), indent=Inches(0.2))

    if amnesty.get("reasons"):
        _heading(doc, "Section 128A — waiver of interest and penalty", 2)
        _para(doc,
              "AVAILABLE on these facts, subject to the conditions below."
              if amnesty.get("eligible") else
              "NOT available on these facts.",
              bold=True)
        for reason in amnesty["reasons"]:
            _para(doc, f"•  {reason}", size=Pt(9), indent=Inches(0.2))
        for caveat in amnesty.get("caveats") or []:
            _para(doc, f"•  {caveat}", size=Pt(9), indent=Inches(0.2))

    if limitation.get("computed"):
        _heading(doc, "Limitation for appeal — Section 107", 2)
        _para(doc, limitation.get("message", ""),
              bold=limitation.get("status") in ("condonable", "time_barred"))
        _particulars(doc, [
            ("Order communicated", _date(limitation.get("order_date"))),
            ("Ordinary deadline", _date(limitation.get("ordinary_deadline"))),
            ("Condonable until", _date(limitation.get("condonable_deadline"))),
        ])
        for caveat in limitation.get("caveats") or []:
            _para(doc, f"•  {caveat}", size=Pt(9), indent=Inches(0.2))

    if deposit.get("computed"):
        _heading(doc, "Pre-deposit to maintain an appeal", 2)
        _para(doc, f"{_rupees(deposit.get('amount'))} — "
                   f"{deposit.get('working', '')}")
        _para(doc, deposit.get("basis", ""), size=Pt(9), italic=True)
        for caveat in deposit.get("caveats") or []:
            _para(doc, f"•  {caveat}", size=Pt(9), indent=Inches(0.2))


def _hearing_brief_section(doc: Document, matter: Dict[str, Any],
                           defect_list: List[Dict[str, Any]]):
    """
    What to expect across the table, limb by limb.

    The written reply is half the proceeding. Section 75(4) gives a right of
    hearing wherever an adverse decision is contemplated, and the hearing is
    where a limb answered perfectly on paper gets conceded by an unprepared
    answer. For each limb: the position in one line, the artefact to carry,
    and the questions the officer actually asks on that defect type.
    """
    if not defect_list:
        return

    from .domains import get_pack

    try:
        pack = get_pack(matter.get("domain") or "gst")
    except Exception:
        return
    if not hasattr(pack, "hearing_questions_for"):
        return

    doc.add_page_break()
    _heading(doc, "10.  Personal Hearing — Brief for the Representative", 1)
    _para(doc,
          "For the firm's representative. Take the evidence column with you: "
          "on the reference matter behind this product, a limb was lost not on "
          "law but because one system report was not produced for "
          "verification.",
          italic=True, size=Pt(9))

    for defect in defect_list:
        index = defect.get("index", "")
        heading = defect.get("heading") or "Unheaded limb"
        _heading(doc, f"Limb {index} — {heading}", 2)

        posture = defect.get("posture") or "undecided"
        _particulars(doc, [
            ("Our position", defects.POSTURE_LABEL.get(posture, posture)),
            ("Amount", UNREAD_LABEL + (
                f" — {defect['amount_note']}" if defect.get("amount_note")
                else " — take it from the annexure before the hearing")
             if defect.get("amount_unread")
             else _rupees(defects.defect_total(defect))),
            ("In one line", (defect.get("our_position") or "").strip()[:400]),
        ])

        carry = list(defect.get("annexures") or []) or \
            list(defect.get("evidence_required") or [])
        if carry:
            _para(doc, "Carry to the hearing:", bold=True, size=Pt(9))
            for item in carry[:6]:
                _para(doc, f"•  {item}", size=Pt(9), indent=Inches(0.2))

        gaps = defect.get("evidence_gap") or []
        if gaps:
            _para(doc, "NOT ON RECORD — do not assert these without the "
                       "document:", bold=True, size=Pt(9))
            for gap in gaps[:6]:
                _para(doc, f"•  {gap}", size=Pt(9), indent=Inches(0.2))

        questions = pack.hearing_questions_for(defect.get("type") or "other")
        _para(doc, "Expect to be asked:", bold=True, size=Pt(9))
        for question in questions:
            _para(doc, f"—  {question}", size=Pt(9), indent=Inches(0.2))


# ---------------------------------------------------------------------------
# Backwards compatibility
# ---------------------------------------------------------------------------


def build_reply_pack(matter: Dict[str, Any]) -> bytes:
    """
    Deprecated. Returns the FILING document only.

    Kept so existing callers do not silently break, but it no longer produces
    the combined pack — that combination was the defect this module was
    rewritten to remove. Callers wanting the analysis must ask for the file
    note explicitly, which is the point.
    """
    return build_filing_reply(matter)


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------


def _rupees(value) -> Optional[str]:
    if value in (None, ""):
        return None
    try:
        return f"Rs. {defects.indian_number(float(value))}"
    except (TypeError, ValueError):
        return str(value)


def _date(value) -> Optional[str]:
    if not value:
        return None
    try:
        return datetime.fromisoformat(str(value)).strftime("%d %B %Y")
    except ValueError:
        return str(value)


# Control characters (notably a bare \n or \r) and quote/backslash, which
# terminate or escape a quoted Content-Disposition header value early. A
# client name read off a scanned notice can carry an embedded newline from
# OCR line-wrapping; left in the filename it turns one header into two lines
# and the ASGI server aborts the connection outright — the browser reports
# "Failed to fetch" and the download never starts, for a document that
# otherwise built cleanly.
_FILENAME_UNSAFE_RE = re.compile(r'[\x00-\x1f\x7f"\\]+')


def _clean_filename_part(text: str) -> str:
    return re.sub(r"\s+", " ", _FILENAME_UNSAFE_RE.sub(" ", text)).strip()


def _base_filename(matter: Dict[str, Any]) -> str:
    intake = matter.get("intake", {})
    parts = [
        _clean_filename_part(intake.get("client_name") or "Matter")
            .replace(" ", "_")[:40],
        _clean_filename_part(intake.get("notice_type") or "Notice")
            .replace(" ", ""),
        _clean_filename_part(intake.get("tax_period") or "")
            .replace(" ", "").replace("/", "-")[:20],
    ]
    return "_".join(p for p in parts if p)


def suggested_filename(matter: Dict[str, Any]) -> str:
    """Filename for the filing document."""
    return f"{_base_filename(matter)}_Reply.docx"


def file_note_filename(matter: Dict[str, Any]) -> str:
    """Filename for the internal working paper."""
    return f"{_base_filename(matter)}_File_Note_INTERNAL.docx"
