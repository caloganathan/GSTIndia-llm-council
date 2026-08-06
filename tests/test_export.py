"""Export tests — two documents, and the wall between them.

Three things are locked in here.

CONFIDENTIALITY. The filing document goes to the proper officer. The file note
stays in the office. Every assertion in `TestTheWallBetweenDocuments` exists
because a single combined pack was once produced that carried the firm's own
exposure arithmetic and its own doubts about its case into the document
intended for the department. That must never be possible again, so it is
tested from both sides: the internal content must be present in the file note
AND absent from the filing document.

VERIFICATION GATING. An authority that did not verify may appear in the file
note with a flag. It may not appear in the filing document at all.

TYPOGRAPHY AND REGISTER. Arial 11pt, black only, and no trace in either
document of the machinery that produced it.
"""

import io
import re

import pytest
from docx import Document

from backend import config, export

MATTER = {
    "id": "a1b2c3d4-1111-2222-3333-444455556666",
    "domain": "gst",
    "intake": {
        "client_name": "Acme Industries Private Limited",
        "gstin": "29AAAPL1234C1ZV",
        "notice_type": "ASMT-10",
        "notice_reference": "ZD290626123456X",
        "state": "Karnataka",
        "tax_period": "FY 2019-20",
        "section_invoked": "61",
        "issuing_officer": "R Kumar, Assistant Commissioner",
        "jurisdiction_office": "Jayanagar Assessment Circle, Bengaluru",
        "notice_date": "2026-06-14",
        "due_date": "2026-07-30",
    },
    "metadata": {
        "tier_label": "Pro Council",
        "anonymised": False,
        "models": {"chairman": "some/model", "revenue": "another/model"},
    },
    "result": {
        "determination": {
            "recommended_position": "The notice is barred by limitation.",
            "confidence": "defensible",
            "lead_argument": "Limitation under section 73(10).",
            "preliminary_submissions":
                "1.  The Noticee submits this reply to the notice dated "
                "14 June 2026.\n\n2.  At the outset it is respectfully "
                "submitted that the proceedings are barred by limitation.",
            "defects": [
                {
                    "index": 1,
                    "heading": "Excess input tax credit against GSTR-2B",
                    "type": "itc_excess_2b",
                    "posture": "contested",
                    "strength": "strong",
                    "amount_by_head": {"cgst": 2060000.0, "sgst": 2060000.0},
                    "department_contention": "Excess credit of Rs. 41,20,000.",
                    "our_position": "The conditions in Section 16(2) stand satisfied.",
                    "facts": "1.  The credit was availed against valid invoices.",
                    "legal_framework": [
                        {"provision": "Section 16(2), CGST Act, 2017",
                         "relevance": "Cumulative conditions for eligibility"},
                        {"provision": "Circular No. 999/9/2024-GST",
                         "relevance": "A circular that does not exist"},
                    ],
                    "authorities": [
                        {"citation": "Section 73(10)", "proposition": "Limitation"},
                        {"citation": "Circular No. 183/15/2022-GST",
                         "proposition": "2A mismatch"},
                    ],
                    "submission": "The demand of Rs. 41,20,000 is not sustainable.",
                    "evidence_required": ["Month-wise GSTR-2B"],
                    "evidence_gap": ["Electronic credit ledger for FY 2019-20"],
                    "annexures": ["Month-wise GSTR-2B for FY 2019-20"],
                    "prayer_relief": "DROP the demand of Rs. 41,20,000.",
                },
                {
                    "index": 2,
                    "heading": "GSTR-1 late fee",
                    "type": "late_fee",
                    "posture": "agreed_paid",
                    "amount_by_head": {"cgst": 1150.0, "sgst": 1150.0},
                    "our_position": "Agreed and discharged.",
                    "submission": "The late fee of Rs. 2,300 has been paid.",
                    "payment": {"reference": "AD290626001122B",
                                "date": "26/06/2026", "under_protest": False},
                    "prayer_relief": "ACKNOWLEDGE the payment of Rs. 2,300.",
                },
            ],
            "triage": {"total_count": 2, "argue_count": 1, "settle_count": 1,
                       "total_amount": 4122300.0, "argued_amount": 4120000.0,
                       "settled_amount": 2300.0},
            "filing_blockers": ["Issue 2: no annexure is listed."],
            "risk_flags": ["Penalty exposure under section 122(2)(a)."],
            "documents_to_collect": ["Supplier-wise reconciliation"],
            "panel_disagreements": [{
                "question": "Whether to concede part of the demand",
                "positions": "A partial concession was considered",
                "resolution": "Not adopted.",
            }],
            "board_summary": "Gross exposure of Rs. 45.20 lakh.",
            "working_note": "The question of limitation was taken first.",
            "open_questions": [],
        },
        "verification": {
            "checked": True,
            "summary": {"verified": 2, "unverified": 1, "not_found": 1, "total": 4},
            "authorities": [
                {"citation": "Section 73(10)", "proposition": "Limitation",
                 "status": "VERIFIED", "note": "Traced.", "correction": "",
                 "defect_index": 1},
                {"citation": "Circular No. 183/15/2022-GST",
                 "proposition": "2A mismatch", "status": "UNVERIFIED",
                 "note": "Not confirmed.", "correction": "", "defect_index": 1},
                {"citation": "Section 16(2), CGST Act, 2017",
                 "proposition": "Cumulative conditions", "status": "VERIFIED",
                 "note": "Traced.", "correction": "", "defect_index": 1,
                 "source": "legal_framework"},
                {"citation": "Circular No. 999/9/2024-GST",
                 "proposition": "A circular that does not exist",
                 "status": "NOT_FOUND", "note": "No such circular.",
                 "correction": "", "defect_index": 1,
                 "source": "legal_framework"},
            ],
        },
    },
}


def _filing(matter=None):
    return Document(io.BytesIO(export.build_filing_reply(matter or MATTER)))


def _note(matter=None):
    return Document(io.BytesIO(export.build_file_note(matter or MATTER)))


def _runs(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                yield run
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip():
                            yield run


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        for container in (section.header, section.footer):
            parts.extend(p.text for p in container.paragraphs)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# The wall
# ---------------------------------------------------------------------------


class TestTheWallBetweenDocuments:
    """
    Internal content belongs in the file note and NOWHERE else.

    A combined pack once carried "worst realistic monetary exposure … approx
    Rs. 2.02 lakh" and "the team has not yet verified line-by-line" in the same
    file as the text meant for the proper officer. Each assertion below is one
    class of that leak.
    """

    INTERNAL_MARKERS = [
        "Gross exposure",                      # board summary
        "Penalty exposure under section 122",  # risk flag
        "Supplier-wise reconciliation",        # documents to collect
        "The question of limitation was taken first",   # working note
        "Whether to concede part of the demand",        # panel disagreement
        "Electronic credit ledger for FY 2019-20",      # evidence gap
        "no annexure is listed",                        # filing blocker
    ]

    @pytest.mark.parametrize("marker", INTERNAL_MARKERS)
    def test_internal_content_is_absent_from_the_filing_document(self, marker):
        assert marker not in _all_text(_filing())

    @pytest.mark.parametrize("marker", INTERNAL_MARKERS)
    def test_internal_content_is_present_in_the_file_note(self, marker):
        assert marker in _all_text(_note())

    def test_the_firms_own_confidence_never_reaches_the_officer(self):
        text = _all_text(_filing())
        for word in ("Defensible", "Weak", "insufficient_information",
                     "Position Recommended"):
            assert word not in text

    def test_file_note_is_stamped_on_every_page(self):
        header = _note().sections[0].header.paragraphs[0].text
        assert "NOT FOR SUBMISSION" in header

    def test_filing_document_is_not_stamped_internal(self):
        assert "NOT FOR SUBMISSION" not in _all_text(_filing())

    def test_the_two_documents_have_different_filenames(self):
        assert export.suggested_filename(MATTER) != export.file_note_filename(MATTER)
        assert "INTERNAL" in export.file_note_filename(MATTER)
        assert "INTERNAL" not in export.suggested_filename(MATTER)


# ---------------------------------------------------------------------------
# Verification gating
# ---------------------------------------------------------------------------


class TestVerificationGating:
    def test_verified_authority_reaches_the_filing_document(self):
        assert "Section 73(10)" in _all_text(_filing())

    def test_unverified_authority_does_not(self):
        assert "Circular No. 183/15/2022-GST" not in _all_text(_filing())

    def test_unverified_authority_is_flagged_in_the_file_note(self):
        text = _all_text(_note())
        assert "Circular No. 183/15/2022-GST" in text
        assert "To be confirmed" in text
        assert "WITHHELD from the filing document" in text

    def test_nothing_is_filable_when_verification_did_not_run(self):
        """
        No verification means no gate passed. Fewer authorities in a reply is
        recoverable; a fabricated one is not.
        """
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["verification"] = {}
        assert "Section 73(10)" not in _all_text(_filing(matter))

    def test_superseded_authority_is_withheld_and_flagged(self):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["verification"]["authorities"][0].update(
            {"status": "SUPERSEDED", "correction": "Circular No. 220/2025"})
        matter["result"]["verification"]["summary"] = {
            "verified": 0, "superseded": 1, "unverified": 1, "total": 2}
        assert "Section 73(10)" not in _all_text(_filing(matter))
        note = _all_text(_note(matter))
        assert "Superseded" in note
        assert "Circular No. 220/2025" in note

    def test_internal_status_enums_never_surface(self):
        for text in (_all_text(_filing()), _all_text(_note())):
            assert "UNVERIFIED" not in text
            assert "NOT_FOUND" not in text

    def test_verified_framework_entry_reaches_the_filing_document(self):
        assert "Cumulative conditions for eligibility" in _all_text(_filing())

    def test_unverified_framework_entry_is_withheld_from_the_filing_document(self):
        """The legal framework table is gated exactly as the authorities
        table is: the chairman is invited to put circulars and case law in
        it, and an entry that failed verification must not print for the
        officer."""
        text = _all_text(_filing())
        assert "Circular No. 999/9/2024-GST" not in text
        assert "Circular No. 999/9/2024-GST" in _all_text(_note())

    def test_no_framework_is_filable_when_verification_did_not_run(self):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["verification"] = {}
        assert "Cumulative conditions for eligibility" not in _all_text(_filing(matter))


# ---------------------------------------------------------------------------
# Citations inside the filed prose
# ---------------------------------------------------------------------------


class TestProseCitationGate:
    """
    A citation the chairman writes INTO a submission paragraph cannot be
    withheld the way a table entry is — the paragraph would have to be
    rewritten. So it is stamped and blocked instead: the filing document
    carries a NOT FOR FILING header until the citation is confirmed or
    struck, and the file note leads with it as a blocker.
    """

    def _with_prose_citation(self, citation_sentence):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["determination"]["defects"][0]["submission"] = (
            f"The demand of Rs. 41,20,000 is not sustainable "
            f"{citation_sentence}"
        )
        return matter

    def test_an_unverified_prose_citation_stamps_the_filing_document(self):
        matter = self._with_prose_citation(
            "in view of Circular No. 183/15/2022-GST")
        header = _filing(matter).sections[0].header.paragraphs[0].text
        assert "NOT FOR FILING" in header
        assert "UNVERIFIED AUTHORITY" in header

    def test_an_unverified_prose_citation_leads_the_file_note(self):
        matter = self._with_prose_citation(
            "in view of Circular No. 183/15/2022-GST")
        note = _all_text(_note(matter))
        assert "cannot be withheld automatically" in note

    def test_a_verified_prose_citation_does_not_stamp(self):
        import copy
        matter = self._with_prose_citation(
            "in view of Circular No. 172/04/2022-GST")
        matter["result"]["verification"]["authorities"].append(
            {"citation": "Circular No. 172/04/2022-GST",
             "proposition": "Works contract credit", "status": "VERIFIED",
             "note": "Traced.", "correction": "", "defect_index": 1,
             "source": "filed_text"})
        assert "NOT FOR FILING" not in _all_text(_filing(matter))

    def test_prose_citations_are_stamped_when_verification_never_ran(self):
        """Same failure direction as _is_filable: a citation nobody checked
        is a gap, not a pass."""
        import copy
        matter = self._with_prose_citation(
            "in view of Circular No. 172/04/2022-GST")
        matter["result"]["verification"] = {}
        header = _filing(matter).sections[0].header.paragraphs[0].text
        assert "NOT FOR FILING" in header

    def test_clean_prose_is_not_stamped(self):
        assert "NOT FOR FILING" not in _all_text(_filing())

    def test_a_blocker_the_panel_already_recorded_is_not_listed_twice(self):
        matter = self._with_prose_citation(
            "in view of Circular No. 183/15/2022-GST")
        matter["result"]["determination"]["filing_blockers"] = [
            "The reply text for defect 1 cites an authority that did not "
            "verify: Circular No. 183/15/2022-GST [UNVERIFIED]. Confirm it "
            "against the reported text or remove it from the prose."
        ]
        note = _all_text(_note(matter))
        assert "1 matter(s) must be resolved before filing" in note

    def test_the_stamp_joins_an_existing_watermark_rather_than_replacing_it(self):
        import copy
        matter = self._with_prose_citation(
            "in view of Circular No. 183/15/2022-GST")
        matter["metadata"]["watermark"] = "DRAFT — NOT FOR FILING"
        header = _filing(matter).sections[0].header.paragraphs[0].text
        assert "DRAFT" in header
        assert "UNVERIFIED AUTHORITY" in header


# ---------------------------------------------------------------------------
# The filing document is a filing document
# ---------------------------------------------------------------------------


class TestFilingDocumentIsFilable:
    def test_carries_a_cause_title(self):
        text = _all_text(_filing())
        assert "BEFORE THE ASSISTANT COMMISSIONER" in text
        assert "JAYANAGAR ASSESSMENT CIRCLE, BENGALURU" in text
        assert "GOVERNMENT OF KARNATAKA" in text

    def test_carries_the_notice_reference(self):
        """A reply that does not quote the reference can be rejected on sight."""
        assert "ZD290626123456X" in _all_text(_filing())

    def test_carries_a_signature_block(self):
        text = _all_text(_filing())
        assert "Yours faithfully" in text
        assert "Authorised Signatory" in text
        assert "For M/s. Acme Industries Private Limited" in text

    def test_signature_block_names_the_client_not_the_firm(self, monkeypatch):
        """Replies go out on the client's letterhead unless asked otherwise."""
        monkeypatch.setattr(config, "FIRM_NAME", "JCSS & Associates LLP")
        assert "JCSS & Associates LLP" not in _all_text(_filing())

    def test_has_the_a_to_o_sections(self):
        text = _all_text(_filing())
        for section in ("A.  PRELIMINARY SUBMISSIONS",
                        "B.  DISPUTES AT A GLANCE",
                        "C.  ISSUE-WISE DETAILED REPLY",
                        "D.  CONSOLIDATED STATEMENT OF PAYMENTS",
                        "E.  EVIDENTIARY INDEX",
                        "F.  PRAYER"):
            assert section in text.upper(), f"missing section: {section}"

    def test_prayer_has_one_relief_per_defect(self):
        text = _all_text(_filing())
        assert "DROP the demand of Rs. 41,20,000." in text
        assert "ACKNOWLEDGE the payment of Rs. 2,300." in text
        assert "And for this act of justice" in text

    def test_prayer_always_seeks_a_hearing(self):
        assert "GRANT a personal hearing" in _all_text(_filing())

    def test_each_defect_states_its_own_figures(self):
        text = _all_text(_filing())
        assert "Rs. 41,20,000" in text
        assert "CGST 20,60,000 + SGST 20,60,000" in text

    def test_payment_reference_is_stated(self):
        """An officer closes a conceded limb on the DRC-03 reference."""
        assert "AD290626001122B" in _all_text(_filing())

    def test_annexures_are_numbered_and_mapped(self):
        text = _all_text(_filing())
        assert "Annexure-1" in text
        assert "Month-wise GSTR-2B for FY 2019-20" in text

    def test_under_protest_reservation_is_stated(self):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["determination"]["defects"][1]["posture"] = \
            "paid_under_protest"
        matter["result"]["determination"]["defects"][1]["payment"][
            "under_protest"] = True
        text = _all_text(_filing(matter))
        assert "without prejudice" in text
        assert "reserves its right to seek refund" in text

    def test_watermark_stamped_when_the_tier_requires_it(self):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["metadata"]["watermark"] = "DRAFT — NOT FOR FILING"
        assert "DRAFT — NOT FOR FILING" in _all_text(_filing(matter))


# ---------------------------------------------------------------------------
# Typography and register
# ---------------------------------------------------------------------------


class TestTypography:
    @pytest.mark.parametrize("builder", [_filing, _note])
    def test_arial_throughout(self, builder):
        fonts = {r.font.name for r in _runs(builder())}
        assert fonts <= {"Arial", None}, f"non-Arial fonts present: {fonts}"

    @pytest.mark.parametrize("builder", [_filing, _note])
    def test_body_is_eleven_point(self, builder):
        assert builder().styles["Normal"].font.size.pt == 11

    @pytest.mark.parametrize("builder", [_filing, _note])
    def test_monochrome_only(self, builder):
        colours = {
            str(r.font.color.rgb)
            for r in _runs(builder())
            if r.font.color and r.font.color.rgb
        }
        assert colours <= {"000000", "404040"}, f"colour present: {colours}"

    def test_section_headings_are_bold(self):
        doc = _filing()
        headings = [
            p for p in doc.paragraphs
            if p.runs and p.runs[0].bold and re.match(r"^[A-F]\.\s", p.text.strip())
        ]
        assert len(headings) >= 5

    @pytest.mark.parametrize("builder", [_filing, _note])
    def test_no_shading_applied(self, builder):
        xml = builder().element.xml
        assert 'w:fill="' not in xml or 'w:fill="auto"' in xml or \
               xml.count('w:fill="') == xml.count('w:fill="auto"')


class TestNoMachineFingerprints:
    BANNED = [
        r"\bAI\b", r"\bLLM\b", "artificial intelligence", "automated",
        r"\bmodel\b", r"\bpanel\b", r"\bchairman\b", r"\bcounsel\b",
        r"\bcouncil\b", "advocate", "machine.generated", "deliberat",
        "openrouter", r"\bgpt\b", r"\bclaude\b", r"\bgemini\b", r"\bgrok\b",
        "deepseek", "prompt",
    ]

    @pytest.mark.parametrize("pattern", BANNED)
    def test_filing_document_carries_no_machinery(self, pattern):
        text = _all_text(_filing())
        assert not re.search(pattern, text, re.IGNORECASE), \
            f"machine vocabulary leaked into the filing document: {pattern}"

    def test_provenance_omitted_by_default(self):
        text = _all_text(_note()).lower()
        assert "some/model" not in text

    def test_provenance_available_for_the_firms_own_file(self, monkeypatch):
        monkeypatch.setattr(config, "EXPORT_PROVENANCE", True)
        text = _all_text(_note())
        assert "some/model" in text

    def test_provenance_never_reaches_the_filing_document(self, monkeypatch):
        monkeypatch.setattr(config, "EXPORT_PROVENANCE", True)
        assert "some/model" not in _all_text(_filing())


# ---------------------------------------------------------------------------
# The file note
# ---------------------------------------------------------------------------


class TestFileNote:
    def test_leads_with_what_blocks_filing(self):
        text = _all_text(_note())
        assert "Before this reply can be filed".upper() in text.upper()

    def test_evidence_gaps_have_their_own_section(self):
        """
        The one limb lost in the matter this was built against was lost on a
        missing document, not on an argument.
        """
        text = _all_text(_note())
        assert "EVIDENCE GAPS" in text.upper()
        assert "Electronic credit ledger for FY 2019-20" in text

    def test_records_the_triage_split(self):
        text = _all_text(_note())
        assert "2 defects" in text
        assert "1 argued" in text

    def test_defect_register_lists_every_limb(self):
        text = _all_text(_note())
        assert "Excess input tax credit against GSTR-2B" in text
        assert "GSTR-1 late fee" in text

    def test_review_note_present(self):
        assert "settled and signed by the engagement partner" in _all_text(_note())

    def test_firm_name_printed_when_configured(self, monkeypatch):
        monkeypatch.setattr(config, "FIRM_NAME", "JCSS & Associates LLP")
        assert "JCSS & Associates LLP" in _all_text(_note())

    def test_no_firm_heading_when_unset(self, monkeypatch):
        monkeypatch.setattr(config, "FIRM_NAME", "")
        assert "Chartered Accountants" not in _all_text(_note())

    def test_unparseable_chairman_output_is_preserved(self):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["determination"]["unstructured_output"] = \
            "raw text that could not be parsed"
        assert "raw text that could not be parsed" in _all_text(_note(matter))


# ---------------------------------------------------------------------------
# Conventions and degradation
# ---------------------------------------------------------------------------


class TestIndianConventions:
    def test_rupee_digit_grouping(self):
        assert export._rupees(4520000) == "Rs. 45,20,000"
        assert export._rupees(100000) == "Rs. 1,00,000"
        assert export._rupees(999) == "Rs. 999"
        assert export._rupees(12345678) == "Rs. 1,23,45,678"

    def test_rupees_handles_bad_input(self):
        assert export._rupees(None) is None
        assert export._rupees("") is None
        assert export._rupees("a bucket range") == "a bucket range"

    def test_dates_in_long_form(self):
        assert export._date("2026-06-14") == "14 June 2026"
        assert export._date(None) is None
        assert export._date("not a date") == "not a date"

    def test_filenames_are_descriptive(self):
        reply = export.suggested_filename(MATTER)
        assert reply.endswith("_Reply.docx")
        assert "Acme_Industries_Private_Limited" in reply
        assert "ASMT-10" in reply
        assert export.file_note_filename(MATTER).endswith(
            "_File_Note_INTERNAL.docx")


class TestMalformedPanelOutputStillBuilds:
    """
    Neither document may fail to build over one malformed limb.

    A builder that raises does not degrade gracefully — the response dies
    mid-flight and the browser reports "Failed to fetch", with nothing in the
    UI to say which limb was at fault. A limb rendered without its payment row
    is recoverable; a download that never arrives is not.
    """

    def test_payment_written_as_prose_does_not_break_the_filing_reply(self):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["determination"]["defects"][1]["payment"] = \
            "Rs. 2,300 paid vide DRC-03 dated 26/06/2026"
        text = _all_text(_filing(matter))
        assert "GSTR-1 late fee" in text

    def test_payment_written_as_prose_does_not_break_the_file_note(self):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["determination"]["defects"][1]["payment"] = "paid"
        assert "GSTR-1 late fee" in _all_text(_note(matter))

    @pytest.mark.parametrize("builder", [_filing, _note])
    def test_a_non_dict_limb_is_skipped_rather_than_raising(self, builder):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["determination"]["defects"].append("a stray string")
        assert "Excess input tax credit against GSTR-2B" in _all_text(builder(matter))


class TestFilenameSanitization:
    """
    A client name lifted from a scanned notice can carry stray control
    characters — most often an embedded newline from OCR line-wrapping. Left
    in the Content-Disposition header, a raw newline splits it into two
    header lines: uvicorn aborts the connection mid-response and the browser
    reports "Failed to fetch", for a document that otherwise built cleanly.
    """

    def test_embedded_newline_in_client_name_does_not_break_the_filename(self):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["intake"]["client_name"] = "Acme Industries\nPrivate Limited"
        for name in (export.suggested_filename(matter),
                     export.file_note_filename(matter)):
            assert "\n" not in name
            assert "\r" not in name

    def test_quote_and_backslash_are_stripped(self):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["intake"]["client_name"] = 'Acme "Steel" Ind\\ustries'
        name = export.suggested_filename(matter)
        assert '"' not in name
        assert "\\" not in name

    def test_no_control_characters_survive_at_all(self):
        """
        Any control character (not just \\n) can break header framing the
        same way. Sweep the whole range, not just the one that was seen.
        """
        import copy
        matter = copy.deepcopy(MATTER)
        matter["intake"]["client_name"] = "Acme\x00Industries\x1fLtd\x7f"
        name = export.suggested_filename(matter)
        assert not any(ord(c) < 0x20 or ord(c) == 0x7f for c in name)


class TestDegradedInput:
    def test_empty_determination_still_produces_both_documents(self):
        matter = {**MATTER, "result": {"determination": {}, "verification": {}}}
        assert "PRELIMINARY SUBMISSIONS" in _all_text(_filing(matter)).upper()
        assert "POSITION RECOMMENDED" in _all_text(_note(matter)).upper()

    def test_filing_document_falls_back_to_a_default_opening(self):
        matter = {**MATTER, "result": {"determination": {}, "verification": {}}}
        text = _all_text(_filing(matter))
        assert "hereinafter referred to as 'the Noticee'" in text
        assert "29AAAPL1234C1ZV" in text

    def test_missing_authorities_is_called_out_in_the_file_note(self):
        matter = {**MATTER, "result": {
            "determination": {"recommended_position": "x"},
            "verification": {"authorities": []},
        }}
        assert "carries no case law" in _all_text(_note(matter))

    def test_no_blockers_is_stated_positively(self):
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["determination"]["filing_blockers"] = []
        assert "No structural blockers" in _all_text(_note(matter))


class TestDepartmentAllegationIsProse:
    """
    The department's allegation is quoted in the filed reply. When the panel
    does not supply a clean one it falls back to the notice extract, and a
    departmental annexure flattened into text is a paragraph of loose digits
    and column rulers. Pasting that into a document filed with the officer
    reads as though nobody looked at it.
    """

    def test_a_flattened_table_is_dropped(self):
        cleaned = export._clean_contention(
            "Excess claim of ITC availed w.r.t GSTR-2B: "
            "SGST CGST IGST CESS Total 1 2 3 4 5 6 1 ITC as per GSTR-2B 8A of "
            "GSTR-9 3801377 3801377 1470165 0 9072919 2 Net eligible 3801377"
        )
        assert cleaned == "Excess claim of ITC availed w.r.t GSTR-2B:"

    def test_column_headers_are_dropped(self):
        assert "Table No." not in export._clean_contention(
            "Credit notes: S.No Description Table No. The details follow.")

    def test_prose_survives_intact(self):
        prose = ("Section 17(5) of the GST Act 2017 states that input tax "
                 "credit shall not be available in respect of certain "
                 "transactions mentioned therein.")
        assert export._clean_contention(prose) == prose

    def test_a_sentence_quoting_one_figure_survives(self):
        """A figure in prose is not a table."""
        prose = "A short payment of Rs. 90,003 has been identified."
        assert export._clean_contention(prose) == prose

    def test_leading_bullet_is_stripped(self):
        assert export._clean_contention("• Difference in turnover: found.") \
            .startswith("Difference in turnover")

    def test_empty_input_is_safe(self):
        assert export._clean_contention(None) == ""
        assert export._clean_contention("") == ""


class TestUnallocatedAmountDisplay:
    def test_no_redundant_parenthetical_when_heads_are_unknown(self):
        """"Rs. 4,81,548 (4,81,548)" reads as a mistake, because it is one."""
        import copy
        matter = copy.deepcopy(MATTER)
        matter["result"]["determination"]["defects"][0]["amount_by_head"] = \
            {"unallocated": 481548.0}
        text = _all_text(_filing(matter))
        assert "Rs. 4,81,548." in text
        assert "Rs. 4,81,548 (4,81,548)" not in text


class TestAnUnreadAmountIsNeverAZero:
    """
    CLAUDE.md gotcha 7: "Extraction that cannot read a figure must report it
    unread. Never fill it with a zero." Every render site printed
    `defect_total()` regardless, which returns 0 for an unread limb — so the
    defect register showed "0" and the hearing brief "Rs. 0.00", and a
    reviewer reads that as "nothing in issue on this limb". `amount_note`,
    which the chairman is explicitly instructed to fill for exactly this case,
    was rendered nowhere at all. The whole path had no test.
    """

    def _matter(self):
        import copy
        matter = copy.deepcopy(MATTER)
        limb = matter["result"]["determination"]["defects"][1]
        limb["amount_by_head"] = {}
        limb["amount_unread"] = True
        limb["amount_note"] = ("The figure sits in Annexure B, which did not "
                               "extract; take it from the department's table.")
        return matter

    def test_the_file_note_does_not_print_a_zero_for_an_unread_limb(self):
        text = _all_text(_note(self._matter()))
        register_zero = "GSTR-1 late fee\n0" in text or "\t0\t" in text
        assert not register_zero, "an unread limb was rendered as 0"
        assert "Not read from the notice" in text

    def test_the_amount_note_reaches_the_reviewer(self):
        text = _all_text(_note(self._matter()))
        assert "Annexure B" in text, (
            "amount_note is the chairman's explanation of what must be taken "
            "off the annexure, and it was rendered nowhere"
        )

    def test_an_unread_limb_is_a_filing_blocker(self):
        text = _all_text(_note(self._matter()))
        assert "the amount could not be read from the notice" in text
        # Level-1 headings render uppercase.
        assert "BEFORE THIS REPLY CAN BE FILED" in text

    def test_the_hearing_brief_does_not_print_rupees_zero(self):
        text = _all_text(_note(self._matter()))
        assert "Rs. 0.00" not in text

    def test_the_file_note_marks_the_matter_total_incomplete(self):
        text = _all_text(_note(self._matter()))
        assert "INCOMPLETE" in text
        assert "1 limb(s)" in text

    def test_the_filing_document_quotes_the_departments_own_total(self):
        """The limb sum understates when a limb is unread. Quoting the short
        figure to the officer states a wrong number in the taxpayer's favour;
        the notice's own declared total is the department's figure."""
        matter = self._matter()
        matter["intake"]["amount_disputed"] = 4122300.0
        text = _all_text(_filing(matter))
        assert "41,22,300" in text

    def test_the_filing_document_never_admits_the_figure_was_unreadable(self):
        """The officer is not told we could not read their own annexure —
        that is an internal problem and an unnecessary admission."""
        text = _all_text(_filing(self._matter()))
        assert "Not read from the notice" not in text
        assert "INCOMPLETE" not in text

    def test_a_fully_read_matter_is_unaffected(self):
        note, filing = _all_text(_note()), _all_text(_filing())
        assert "Not read from the notice" not in note
        assert "INCOMPLETE" not in note
        assert "INCOMPLETE" not in filing
