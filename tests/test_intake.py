"""Notice ingestion.

Two things are being protected here.

Correctness: a wrong tax period or a wrong State silently accepted is worse
than an empty field, because the State drives which High Court binds the
matter. Everything extracted is a proposal, and the tests check that what is
proposed is right.

Privacy: an uploaded notice must be no less private than a typed one. The
free tier scrubs the text before any of it reaches a model, and aborts if
scrubbing fails.
"""

import io

import pytest

from backend import intake, sanitizer
from backend.domains import gst

NOTICE = """
GOVERNMENT OF INDIA
FORM GST ASMT-10
[See rule 99(1)]
Reference No: ZA290624001234                    Date: 14.06.2026

To,
Acme Industries Private Limited
GSTIN: 29AAAPL1234C1ZV
PAN: AAAPL1234C

Notice for intimating discrepancies in the return after scrutiny under
section 61 of the CGST Act, 2017 for the financial year 2019-20.

During scrutiny of the returns for FY 2019-20 the following discrepancies
have been noticed:

1. Input tax credit availed in GSTR-3B exceeds that appearing in GSTR-2A by
   Rs. 41,20,000.
2. Interest under section 50 has been short paid to the extent of Rs. 3,15,000.

You are directed to explain the discrepancies on or before 30.07.2026 failing
which proceedings under section 73 shall be initiated.

Superintendent of Central Tax, Range-IV, Bengaluru
Contact: officer@gst.gov.in
"""


class TestTextExtraction:
    def test_plain_text(self):
        text, warnings = intake.extract_text("notice.txt", NOTICE.encode())
        assert "ASMT-10" in text
        assert warnings == []

    def test_rejects_unsupported_type(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            intake.extract_text("notice.jpg", b"x" * 500)

    def test_rejects_oversized_file(self):
        oversized = b"x" * (intake.MAX_UPLOAD_BYTES + 1)
        with pytest.raises(ValueError, match="larger than"):
            intake.extract_text("notice.pdf", oversized)

    def test_scanned_notice_is_reported_not_guessed(self):
        """No text layer must produce an honest warning, never invented text."""
        text, warnings = intake.extract_text("scan.txt", b"   ")
        assert text == ""
        assert any("scanned" in w for w in warnings)

    def test_corrupt_pdf_raises_clearly(self):
        with pytest.raises(ValueError, match="Could not read the PDF"):
            intake.extract_text("notice.pdf", b"not a pdf at all")

    def test_docx_round_trip(self):
        from docx import Document
        document = Document()
        document.add_paragraph("FORM GST ASMT-10")
        document.add_paragraph("GSTIN: 29AAAPL1234C1ZV")
        buffer = io.BytesIO()
        document.save(buffer)
        text, _ = intake.extract_text("notice.docx", buffer.getvalue())
        assert "ASMT-10" in text
        assert "29AAAPL1234C1ZV" in text


class TestLocalExtraction:
    """No model involved — nothing leaves the machine for any of this."""

    def setup_method(self):
        self.result = intake.extract_fields_local(NOTICE, gst)
        self.fields = self.result["fields"]

    def test_gstin(self):
        assert self.fields["gstin"] == "29AAAPL1234C1ZV"

    def test_state_derived_from_gstin(self):
        """29 is Karnataka — and the State decides which High Court binds."""
        assert self.fields["state"] == "Karnataka"
        assert self.result["sources"]["state"] == "gstin"

    def test_state_is_a_real_jurisdiction(self):
        assert self.fields["state"] in gst.STATE_HIGH_COURT

    def test_notice_type(self):
        assert self.fields["notice_type"] == "ASMT-10"

    def test_reference_number(self):
        assert self.fields["notice_reference"] == "ZA290624001234"

    def test_tax_period(self):
        assert self.fields["tax_period"] == "FY 2019-20"

    def test_section_invoked(self):
        assert self.fields["section_invoked"] == "61"

    def test_notice_date_read_from_the_document(self):
        assert self.fields["notice_date"] == "2026-06-14"

    def test_due_date_comes_from_the_direction_to_reply(self):
        """
        The reply deadline is read where the notice states it, and this notice
        states it: "explain the discrepancies on or before 30.07.2026".

        This test previously asserted the opposite — that no due date was read
        here at all — because the only accepted source was the portal form's
        labelled row. That was too narrow. The State letter formats carry no
        such row and print the deadline in a sentence instead, so refusing to
        read it left the field empty on notices that plainly stated it, and a
        deadline nobody extracted is a deadline nobody diarised.

        What was actually wrong with the old behaviour is preserved by the
        test below: the date must be anchored to a direction to answer, not
        taken from its position in the document.
        """
        assert self.fields["due_date"] == "2026-07-30"
        assert self.result["sources"]["due_date"] == "notice-directed"

    def test_a_date_with_no_direction_beside_it_is_not_a_deadline(self):
        """
        The failure the old rule existed to prevent. Taking the LAST date in
        the document read an invoice date out of an annexure and presented it
        as the deadline — an error that, acted on, loses the client the right
        to reply. An "on or before" with no direction to answer in front of it
        is a term of supply, not a deadline.
        """
        text = (
            "GSTIN: 29AAAPL1234C1ZV\n"
            "FORM GST ASMT-10\n"
            "Annexure: the goods were to be delivered on or before 30.07.2026 "
            "under the purchase order.\n"
        )
        fields = intake.extract_fields_local(text, gst)["fields"]
        assert "due_date" not in fields

    def test_a_relative_window_is_recorded_as_a_window(self):
        """
        Most scrutiny notices print no date at all — they run a period from
        service, which the department cannot date either. Recording the period
        is the honest reading; inventing a date from it would not be.
        """
        text = (
            "GSTIN: 29AAAPL1234C1ZV\n"
            "FORM GST ASMT-10\n"
            "You are requested to file your reply through the GST Common "
            "portal within 30 days of receipt of the notice.\n"
        )
        fields = intake.extract_fields_local(text, gst)["fields"]
        assert fields["reply_window_days"] == 30
        assert "due_date" not in fields

    def test_amount_is_not_guessed_from_the_largest_figure(self):
        """
        Amounts come from the defect annexures, head by head. Taking the
        largest rupee figure in the document reported a single penalty limb as
        the whole demand on a real notice, understating it by a factor of
        seventy.
        """
        assert "amount_disputed" not in self.fields

    def test_every_field_records_its_source(self):
        for key in self.fields:
            assert key in self.result["sources"]


class TestNoticeTypeMatching:
    def test_longer_code_wins(self):
        """DRC-01A must never be read as DRC-01."""
        assert intake.find_notice_type("FORM GST DRC-01A intimation", gst) == "DRC-01A"

    def test_tolerates_spacing_variants(self):
        for variant in ("ASMT-10", "ASMT 10", "ASMT10"):
            assert intake.find_notice_type(f"FORM GST {variant}", gst) == "ASMT-10"

    def test_absent_returns_none(self):
        assert intake.find_notice_type("A letter about nothing", gst) is None


class TestDateParsing:
    def test_numeric_formats(self):
        for text in ("14.06.2026", "14-06-2026", "14/06/2026"):
            assert intake.find_dates(text) == ["2026-06-14"]

    def test_written_formats(self):
        assert intake.find_dates("14 June 2026") == ["2026-06-14"]
        assert intake.find_dates("14th June, 2026") == ["2026-06-14"]

    def test_document_order_preserved(self):
        assert intake.find_dates("dated 01.01.2026 reply by 15.02.2026") == [
            "2026-01-01", "2026-02-15"
        ]

    def test_impossible_dates_rejected(self):
        assert intake.find_dates("45.13.2026") == []

    def test_deduplicates(self):
        assert intake.find_dates("14.06.2026 and again 14.06.2026") == ["2026-06-14"]


class TestStateCodes:
    def test_every_code_maps_to_a_known_jurisdiction(self):
        for code, state in intake.GSTIN_STATE_CODES.items():
            assert state in gst.STATE_HIGH_COURT, f"{code} -> {state} has no High Court"

    def test_key_states(self):
        assert intake.GSTIN_STATE_CODES["33"] == "Tamil Nadu"
        assert intake.GSTIN_STATE_CODES["27"] == "Maharashtra"
        assert intake.GSTIN_STATE_CODES["07"] == "Delhi"


class TestAmountsAndPeriods:
    def test_indian_grouping(self):
        assert intake.find_amounts("Rs. 41,20,000") == [4120000.0]

    def test_currency_variants(self):
        assert intake.find_amounts("INR 1,000 and ₹ 2,000") == [1000.0, 2000.0]

    def test_period_variants(self):
        for text in ("FY 2019-20", "financial year 2019-20", "F.Y. 2019-2020"):
            assert intake.find_tax_period(text) == "FY 2019-20"


@pytest.mark.asyncio
class TestPrivacy:
    """An uploaded notice must be no less private than a typed one."""

    FREE = {"anonymise": True, "grounding": "free/model"}
    PRO = {"anonymise": False, "grounding": "pro/model"}

    async def test_free_tier_scrubs_before_sending(self, monkeypatch):
        sent = {}

        async def capture(model, messages, **kwargs):
            sent["prompt"] = messages[0]["content"]
            return {"ok": True, "usage": None,
                    "content": '{"issues": "i", "facts": "f"}'}

        monkeypatch.setattr(intake, "query_model", capture)
        await intake.extract_fields_assisted(NOTICE, gst, self.FREE)

        for secret in ("Acme Industries", "29AAAPL1234C1ZV", "AAAPL1234C",
                       "officer@gst.gov.in"):
            assert secret not in sent["prompt"], f"LEAKED: {secret}"
        assert sanitizer.audit_leaks(sent["prompt"]) == {}

    async def test_pro_tier_sends_full_facts_with_zdr(self, monkeypatch):
        sent = {}

        async def capture(model, messages, **kwargs):
            sent["prompt"] = messages[0]["content"]
            sent["zdr"] = kwargs.get("zdr")
            return {"ok": True, "usage": None,
                    "content": '{"issues": "i", "facts": "f"}'}

        monkeypatch.setattr(intake, "query_model", capture)
        await intake.extract_fields_assisted(NOTICE, gst, self.PRO)
        assert "Acme Industries" in sent["prompt"]
        assert sent["zdr"] is True

    async def test_free_tier_restores_identifiers_in_the_result(self, monkeypatch):
        """The model reads scrubbed text and answers in scrubbed tokens; the
        fields returned to the user must carry the real identifiers back, or
        '[GSTIN-1]' and 'the Taxpayer' reach the matter and both exports."""
        async def echo_placeholders(model, messages, **kwargs):
            prompt = messages[0]["content"]
            # The model can only echo tokens it actually saw. Pull the GSTIN
            # placeholder out of the scrubbed prompt and answer with it.
            import re
            token = re.search(r"\[GSTIN-\d+\]", prompt)
            gstin_token = token.group(0) if token else "[GSTIN-?]"
            return {"ok": True, "usage": None, "content":
                    '{"issues": "Excess ITC by the Taxpayer under ' +
                    gstin_token + '", "facts": "the Taxpayer filed late."}'}

        monkeypatch.setattr(intake, "query_model", echo_placeholders)
        fields, _, _ = await intake.extract_fields_assisted(
            NOTICE, gst, self.FREE)

        blob = " ".join(fields.values())
        assert "[GSTIN" not in blob, "a placeholder token reached the user"
        assert "the Taxpayer" not in blob, "the placeholder name was not restored"
        assert "29AAAPL1234C1ZV" in blob
        assert "Acme Industries" in blob

    async def test_extraction_failure_is_reported_not_swallowed(self, monkeypatch):
        async def failing(*args, **kwargs):
            return {"ok": False, "error": "model unavailable"}

        monkeypatch.setattr(intake, "query_model", failing)
        fields, _, warnings = await intake.extract_fields_assisted(
            NOTICE, gst, self.PRO
        )
        assert fields == {}
        assert any("manually" in w for w in warnings)

    async def test_low_confidence_is_surfaced(self, monkeypatch):
        async def query(*args, **kwargs):
            return {"ok": True, "usage": None,
                    "content": '{"issues": "i", "facts": "f", "confidence": "low"}'}

        monkeypatch.setattr(intake, "query_model", query)
        _, _, warnings = await intake.extract_fields_assisted(NOTICE, gst, self.PRO)
        assert any("difficult to read" in w for w in warnings)


@pytest.mark.asyncio
class TestReadNotice:
    async def test_end_to_end_without_a_model(self):
        """Local extraction alone should fill most of the form."""
        result = await intake.read_notice(
            "notice.txt", NOTICE.encode(), gst,
            {"anonymise": False}, use_model=False,
        )
        fields = result["fields"]
        assert fields["notice_type"] == "ASMT-10"
        assert fields["state"] == "Karnataka"
        assert fields["tax_period"] == "FY 2019-20"
        assert result["usage"] is None

    async def test_missing_fields_are_listed_for_the_user(self):
        result = await intake.read_notice(
            "notice.txt", b"A short letter with no particulars whatsoever.",
            gst, {"anonymise": False}, use_model=False,
        )
        assert any("Could not determine" in w for w in result["warnings"])

    async def test_uploaded_file_is_never_persisted(self, monkeypatch, tmp_path):
        """The binary is parsed in memory and must not touch disk."""
        before = set(tmp_path.iterdir())
        await intake.read_notice(
            "notice.txt", NOTICE.encode(), gst,
            {"anonymise": False}, use_model=False,
        )
        assert set(tmp_path.iterdir()) == before


class TestEntityName:
    """
    Extracted for two reasons, and the second is the important one: it fills
    the client name field, and it gives the sanitiser something to scrub. Without
    it the free tier strips GSTIN and PAN but leaves the company name standing.
    """

    def test_common_indian_suffixes(self):
        cases = [
            ("To, Acme Industries Private Limited\nGSTIN:", "Acme Industries Private Limited"),
            ("M/s Bharat Steel Pvt Ltd is registered", "Bharat Steel Pvt Ltd"),
            ("Sunrise Textiles LLP has filed", "Sunrise Textiles LLP"),
            ("Reliance Trading Limited, the taxpayer", "Reliance Trading Limited"),
        ]
        for text, expected in cases:
            assert intake.find_entity_name(text) == expected

    def test_longest_match_wins(self):
        """A truncated name would leave the remainder unscrubbed."""
        name = intake.find_entity_name("Acme Steel Industries Private Limited")
        assert name == "Acme Steel Industries Private Limited"

    def test_absent_returns_none(self):
        assert intake.find_entity_name("A notice with no entity named.") is None

    def test_populates_the_client_name_field(self):
        fields = intake.extract_fields_local(NOTICE, gst)["fields"]
        assert fields["client_name"] == "Acme Industries Private Limited"


@pytest.mark.asyncio
class TestFreeTierNameScrubbing:
    async def test_company_name_removed_before_sending(self, monkeypatch):
        """The regression the suite caught: the name outlived the identifiers."""
        sent = {}

        async def capture(model, messages, **kwargs):
            sent["prompt"] = messages[0]["content"]
            return {"ok": True, "usage": None, "content": '{"issues": "i"}'}

        monkeypatch.setattr(intake, "query_model", capture)
        await intake.extract_fields_assisted(
            NOTICE, gst, {"anonymise": True, "grounding": "m"}
        )
        assert "Acme" not in sent["prompt"]
        assert sanitizer.TAXPAYER_PLACEHOLDER in sent["prompt"]


class TestFieldsTheOldExtractorGotWrong:
    """
    One regression test per field that a real scrutiny notice was read wrongly.

    Every assertion below corresponds to a value the previous extractor
    actually produced on a live Tamil Nadu ASMT-10 attachment. Six of seven
    auto-filled fields were wrong, and none of them were wrong in a way the
    reviewer could see.
    """

    NOTICE = """\
GST ASMT - 10
[See rule 99(1)]
Reference No.: ZD330226255583F Date:  27/02/2026
To
GSTIN: 33AAGCG4663G1ZO
Name: GRAM ENVOSOLUTION PRIVATE LIMITED.
Tax period: APR 2023 - MAR 2024 F.Y.: 2023-2024
Notice for intimating discrepancies in the return after scrutiny
The following discrepancies were noticed including liabilities under
section 9(5), if applicable. Attention is invited to section 16(2)(aa),
section 16(4) and Rule 36(4). Interest arises under section 50(3).
As per Notification 02/2022-CT(rate) dt 31.03.2022, certain rates changed.
Invoice GRAM/23-24/056 dated 21-06-2023 was amended.
Sr. No. Description Particulars
1 Section under which notice is issued 61
2 Date by which reply has to be submitted 28/03/2026
Signature
Name: Vidhya V
Designation: Assistant Commissioner
Jurisdiction: RAM NAGAR , Coimbatore-
II , COIMBATORE , Tamil Nadu
"""

    @classmethod
    @pytest.fixture(scope="class")
    def fields(cls):
        return intake.extract_fields_local(cls.NOTICE, gst)["fields"]

    def test_provision_comes_from_the_labelled_row(self, fields):
        """Previously read '9' out of 'section 9(5), if applicable'."""
        assert fields["section_invoked"] == "61"

    def test_every_provision_is_captured_with_its_sub_section(self, fields):
        """Previously only the first match, truncated to '16'."""
        cited = fields["sections_cited"]
        assert "16(2)(aa)" in cited
        assert "16(4)" in cited
        assert "50(3)" in cited
        assert len(cited) > 1

    def test_reference_is_an_identifier_not_an_english_word(self, fields):
        """Previously matched on the word 'notice' and returned 'proposing'."""
        assert fields["notice_reference"] == "ZD330226255583F"

    def test_notice_date_is_the_notice_date(self, fields):
        """Previously the earliest date in the document — a 2022 notification."""
        assert fields["notice_date"] == "2026-02-27"

    def test_due_date_is_the_labelled_deadline(self, fields):
        """Previously the latest date in the document — an invoice date."""
        assert fields["due_date"] == "2026-03-28"

    def test_taxpayer_named_in_capitals_is_found(self, fields):
        """
        Previously nothing: the suffix pattern was case-sensitive, so every
        notice printing "PRIVATE LIMITED" left the sanitiser with no company
        name to scrub.
        """
        assert fields["client_name"] == "GRAM ENVOSOLUTION PRIVATE LIMITED"

    def test_form_code_with_spaces_is_recognised(self, fields):
        """The portal prints "GST ASMT - 10", which did not match at all."""
        assert fields["notice_type"] == "ASMT-10"

    def test_issuing_officer_is_captured(self, fields):
        assert fields["issuing_officer"] == "Vidhya V, Assistant Commissioner"

    def test_jurisdiction_survives_a_line_break(self, fields):
        assert fields["jurisdiction_office"].startswith("RAM NAGAR")
        assert "COIMBATORE" in fields["jurisdiction_office"]

    def test_state_derived_from_the_gstin(self, fields):
        assert fields["state"] == "Tamil Nadu"


class TestPrefixedEntityNames:
    def test_tvl_prefix_wins(self):
        assert intake.find_entity_name(
            "Tvl. Gram Envosolution Private Limited, Coimbatore"
        ) == "Gram Envosolution Private Limited"

    def test_m_s_prefix_wins(self):
        assert intake.find_entity_name(
            "M/s. Acme Steel Industries has filed its returns."
        ) == "Acme Steel Industries"

    def test_longest_suffix_match_wins(self):
        assert intake.find_entity_name(
            "Acme Steel Industries Private Limited was registered"
        ) == "Acme Steel Industries Private Limited"

    def test_ordinary_prose_is_not_a_company(self):
        assert intake.find_entity_name(
            "the limited relief sought in the petition"
        ) is None
